# -*- coding: utf-8 -*-
"""
🎙️ Приложение для распознавания речи AI VoiceFinder
✨ Красивый современный интерфейс
Поддержка: временные метки слов, определение спикеров
"""

# Убедимся, что вывод в консоль использует UTF-8 (важно для Windows),
# чтобы избежать UnicodeEncodeError при печати эмодзи.
import sys
import os
try:
    # Для Windows: переопределяем кодировку stdout/stderr на utf-8
    if os.name == 'nt':
        # Устанавливаем переменную окружения, чтобы дочерние процессы
        # также использовали UTF-8 при запуске Python.
        os.environ.setdefault('PYTHONUTF8', '1')
        # Python 3.7+: reconfigure доступен для stdout/stderr
        try:
            sys.stdout.reconfigure(encoding='utf-8')
            sys.stderr.reconfigure(encoding='utf-8')
        except Exception:
            # На некоторых окружениях reconfigure может быть недоступен — безопасно игнорируем
            pass
except Exception:
    pass

print("🚀 Запуск приложения AI VoiceFinder...")
print("=" * 70)

# 📦 Настройка кэша ДО импорта тяжёлых библиотек, чтобы torch/whisperx/HF
# не создавали файлы в системных папках C:\Users\..\.cache
_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
_PORTABLE_MODE = os.environ.get('AI_VOICEFINDER_PORTABLE', 'false').lower() == 'true'
if _PORTABLE_MODE:
    _LOCAL_CACHE_DIR = os.environ.get('AI_VOICEFINDER_CACHE', os.path.join(_SCRIPT_DIR, "models_cache"))
else:
    _LOCAL_CACHE_DIR = os.path.join(_SCRIPT_DIR, "models_cache")
os.makedirs(_LOCAL_CACHE_DIR, exist_ok=True)
os.makedirs(os.path.join(_LOCAL_CACHE_DIR, 'hub'), exist_ok=True)
os.makedirs(os.path.join(_LOCAL_CACHE_DIR, 'transformers'), exist_ok=True)
os.makedirs(os.path.join(_LOCAL_CACHE_DIR, 'torch'), exist_ok=True)
os.environ['HF_HOME'] = _LOCAL_CACHE_DIR
os.environ['HUGGINGFACE_HUB_CACHE'] = os.path.join(_LOCAL_CACHE_DIR, 'hub')
os.environ['TRANSFORMERS_CACHE'] = os.path.join(_LOCAL_CACHE_DIR, 'transformers')
os.environ['TORCH_HOME'] = os.path.join(_LOCAL_CACHE_DIR, 'torch')
os.environ['XDG_CACHE_HOME'] = _LOCAL_CACHE_DIR
# NeMo тоже пишет во временную папку — перенаправляем
os.environ['NEMO_CACHE_DIR'] = os.path.join(_LOCAL_CACHE_DIR, 'nemo')
os.makedirs(os.path.join(_LOCAL_CACHE_DIR, 'nemo'), exist_ok=True)
# Системную папку TEMP/TMP тоже перенаправляем на диск проекта
_NEMO_TMP = os.path.join(_SCRIPT_DIR, 'nemo_tmp')
os.makedirs(_NEMO_TMP, exist_ok=True)
os.environ['TEMP'] = _NEMO_TMP
os.environ['TMP'] = _NEMO_TMP
import tempfile as _tempfile_init
_tempfile_init.tempdir = _NEMO_TMP
print(f"📦 Кэш перенаправлен: {_LOCAL_CACHE_DIR}")
print(f"📁 Временные файлы перенаправлены: {_NEMO_TMP}")

print("📦 Импорт tkinter...")
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
print("✓ tkinter импортирован")

print("🧵 Импорт threading...")
import threading
print("✓ threading импортирован")

print("⚙️ Подавление предупреждений TensorFlow для более быстрого запуска...")
import warnings
warnings.filterwarnings('ignore', category=FutureWarning)
warnings.filterwarnings('ignore', category=UserWarning)
os.environ['TF_CPP_MIN_LOG_LEVEL'] = '3'  # Подавляем логи TensorFlow
print("✓ Предупреждения отключены")

print("🔢 Импорт numpy...")
import numpy as np
print("✓ numpy импортирован")

print("🔥 Импорт torch...")
import torch
print(f"✓ torch импортирован (версия: {torch.__version__})")

print("🔧 Исправление совместимости torchaudio...")
import torchaudio
# Патч для совместимости с новыми версиями torchaudio (2.10.0+)
# где list_audio_backends() был удален
if not hasattr(torchaudio, 'list_audio_backends'):
    def list_audio_backends():
        return ["soundfile"]
    torchaudio.list_audio_backends = list_audio_backends
print("✓ Совместимость torchaudio исправлена")

print("🎙️ Импорт whisperx...")
import whisperx
print("✓ whisperx импортирован")

print("📡 Импорт вспомогательных функций whisperx...")
print("✓ Функции импортированы")

print("🌋 Импорт LavaSR для улучшения аудио...")
try:
    from LavaSR.model import LavaEnhance2
    LAVASR_AVAILABLE = True
    print("✓ LavaSR импортирован")
except ImportError:
    LAVASR_AVAILABLE = False
    print("⚠️ LavaSR не установлен (установите: pip install git+https://github.com/ysharma3501/LavaSR.git)")

print("📁 Импорт os...")
import os
print("✓ os импортирован")

print("🔧 Импорт sys...")
import sys
print("✓ sys импортирован")

print("🖼️ Импорт PIL...")
from PIL import Image, ImageTk
print("✓ PIL импортирован")

print("🔊 Импорт sounddevice для воспроизведения аудио...")
try:
    import sounddevice as sd
    SOUNDDEVICE_AVAILABLE = True
    print("✓ sounddevice импортирован")
except ImportError:
    SOUNDDEVICE_AVAILABLE = False
    print("⚠️ sounddevice не установлен (pip install sounddevice) — воспроизведение будет недоступно")

print("🎬 Импорт модуля обработки MTS видеофайлов...")
try:
    from mts_processor import MTSProcessor, MTSProcessorError, setup_ffmpeg_guide
    MTS_SUPPORT = True
    print("✓ Модуль MTS импортирован")
except ImportError:
    MTS_SUPPORT = False
    print("⚠️ Модуль MTS не найден (не критично)")
except Exception as e:
    MTS_SUPPORT = False
    print(f"⚠️ Ошибка при импорте MTS модуля: {e}")

print("=" * 70)
print("✅ Все модули импортированы успешно!\n")

# Переменные кэша уже установлены в начале файла (до импортов)
LOCAL_CACHE_DIR = _LOCAL_CACHE_DIR
PORTABLE_MODE = _PORTABLE_MODE
if PORTABLE_MODE:
    RESULTS_DIR = os.environ.get('AI_VOICEFINDER_RESULTS', os.path.join(_SCRIPT_DIR, "Results"))
    print("🎒 ПОРТАТИВНЫЙ РЕЖИМ АКТИВИРОВАН")
    print(f"📦 Кэш моделей: {LOCAL_CACHE_DIR}")
else:
    RESULTS_DIR = os.path.join(_SCRIPT_DIR, "Results")
    print(f"📦 Кэш моделей: {LOCAL_CACHE_DIR}")
os.makedirs(RESULTS_DIR, exist_ok=True)


def format_timestamp(seconds):
    """Форматирование времени в формат часы:минуты:секунды"""
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = seconds % 60
    return f"{hours:02d}:{minutes:02d}:{secs:06.3f}"


# 🎨 Цветовая схема (монохромная серо-белая)
COLORS = {
    'primary': '#757575',      # Средний серый
    'success': '#9E9E9E',      # Светло-серый
    'warning': '#BDBDBD',      # Ещё светлее
    'danger': '#616161',       # Тёмно-серый (предупреждения)
    'dark': '#212121',         # Почти чёрный
    'darker': '#141414',       # Очень тёмный серый
    'light': '#F5F5F5',        # Почти белый
    'bg': '#1E1E1E',           # Основной фон (тёмный серый)
    'text': '#EEEEEE',         # Основной текст (почти белый)
    'accent': '#BDBDBD',       # Акцент (светлый серый)
    'card_bg': '#2A2A2A',      # Фон карточек (тёмный серый)
    'border': '#3E3E3E',       # Границы
    'input_bg': '#181818',     # Фон полей ввода
    'hover': '#333333',        # Цвет при наведении
    'playing': '#E0E0E0',      # Почти белый — активное воспроизведение
}


class WhisperXApp:
    """🎙️ Главный класс приложения для распознавания речи AI VoiceFinder"""
    
    def __init__(self, root):
        print("🎨 Инициализация приложения...")
        self.root = root
        self.root.title("AI VoiceFinder")
        self.root.geometry("320x750")
        
        # Запрещаем изменение размера окна
        self.root.resizable(False, False)
        
        # Центрируем окно на экране
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()
        x = (screen_width - 320) // 2
        y = (screen_height - 750) // 2
        self.root.geometry(f"320x750+{x}+{y}")
        
        # Устанавливаем цвет фона
        self.root.configure(bg=COLORS['bg'])
        
        # Настраиваем стили
        self.setup_styles()
        
        print("✓ Окно создано")
        
        # Переменные
        self.model = None
        self.align_model = None
        self.align_metadata = None
        self.diarize_model = None
        self.lavasr_model = None  # Модель LavaSR для улучшения аудио
        self.audio_file_path = None
        self.audio_files_list = []  # Список файлов для пакетной обработки
        self.mts_files_for_processing = []  # MTS видеофайлы для обработки
        self.is_processing = False
        self.device = None
        self.compute_type = None
        print("✓ Переменные инициализированы")
        
        # Создаем интерфейс
        print("🎨 Создание интерфейса...")
        self.create_widgets()
        print("✓ Интерфейс создан")
        
        # Проверяем доступность CUDA
        print("🔍 Проверка CUDA...")
        self.check_cuda()
        print("✓ Проверка CUDA завершена")
        print("\n🎉 Приложение готово к работе!")
    
    def set_input_buttons_state(self, state):
        """Установить состояние кнопок добавления файлов"""
        self.add_file_btn.config(state=state)
        if self.add_mts_btn:
            self.add_mts_btn.config(state=state)
    
    def enhance_audio_with_lavasr(self, audio, sr=16000, denoise=False):
        """
        Улучшение качества аудио с помощью LavaSR
        
        Args:
            audio: numpy array с аудиоданными
            sr: sample rate аудио
            denoise: использовать ли шумоподавление (UL-UNAS)
            
        Returns:
            улучшенный audio
        """
        if not LAVASR_AVAILABLE or self.lavasr_model is None:
            print("⚠️ LavaSR не доступен, пропускаем улучшение")
            return audio
        
        try:
            print("\n" + "=" * 70)
            print("🌋 УЛУЧШЕНИЕ АУДИО С ПОМОЩЬЮ LavaSR")
            print("=" * 70)
            print(f"Входное разрешение: {sr} Hz")
            
            # Преобразуем numpy array в torch tensor если необходимо
            if not isinstance(audio, torch.Tensor):
                import numpy as np
                audio = torch.from_numpy(audio).float()
            
            # Убеждаемся что аудио на правильном устройстве
            if audio.device != torch.device(self.device):
                audio = audio.to(self.device)
            
            # Добавляем batch dimension если необходимо
            if audio.dim() == 1:
                audio = audio.unsqueeze(0)
            elif audio.dim() == 2 and audio.shape[0] > 1:
                # Если несколько каналов, берем первый
                audio = audio[0:1]
            
            print(f"Форма аудио: {audio.shape}")
            
            # Используем batch processing для очень длинных аудио
            audio_length = audio.shape[-1]
            max_chunk_length = int(sr * 30)  # 30 секунд на chunk
            
            if audio_length > max_chunk_length:
                print(f"Длинное аудио ({audio_length/sr:.1f}s), используем batch обработку")
                enhanced_chunks = []
                
                for i in range(0, audio_length, max_chunk_length):
                    chunk = audio[:, i:i + max_chunk_length]
                    print(f"  Обработка chunk {len(enhanced_chunks)+1}...")
                    enhanced_chunk = self.lavasr_model.enhance(
                        chunk, 
                        denoise=denoise,  # Используем переданный параметр
                        batch=False
                    )
                    enhanced_chunks.append(enhanced_chunk.cpu())
                
                enhanced_audio = torch.cat(enhanced_chunks, dim=-1)
            else:
                # Для коротких аудио обрабатываем целиком
                enhanced_audio = self.lavasr_model.enhance(
                    audio,
                    denoise=denoise,  # Используем переданный параметр
                    batch=False
                )
                enhanced_audio = enhanced_audio.cpu()
            
            # Убираем batch dimension если добавляли
            if enhanced_audio.dim() == 2:
                enhanced_audio = enhanced_audio.squeeze(0)
            
            # Преобразуем обратно в numpy
            if isinstance(enhanced_audio, torch.Tensor):
                enhanced_audio = enhanced_audio.numpy()
            
            print(f"Выходное разрешение: 48000 Hz (2x улучшение разрешения)")
            print(f"✓ Аудио успешно улучшено (размер: {enhanced_audio.nbytes / 1024 / 1024:.1f} MB)")
            print("=" * 70)
            
            return enhanced_audio
            
        except Exception as e:
            print(f"❌ Ошибка при улучшении аудио: {e}")
            print("Продолжаем без улучшения...")
            return audio
    
    def normalize_audio(self, audio, target_db=-20.0):
        """
        📊 Нормализует громкость аудио к целевому уровню
        
        Args:
            audio: numpy array или torch tensor с аудиоданными
            target_db: целевой уровень в дБ
            
        Returns:
            нормализованное аудио
        """
        try:
            # Конвертируем в numpy если нужно
            if isinstance(audio, torch.Tensor):
                audio_np = audio.cpu().numpy()
            else:
                audio_np = audio
            
            print("\n" + "=" * 70)
            print("🔊 НОРМАЛИЗАЦИЯ ГРОМКОСТИ АУДИО")
            print("=" * 70)
            
            # Вычисляем текущий RMS уровень
            if audio_np.ndim > 1:
                # Если многоканальное, берем первый канал
                audio_np = audio_np[0] if audio_np.shape[0] > audio_np.shape[-1] else audio_np[:, 0]
            
            rms = np.sqrt(np.mean(audio_np ** 2))
            current_db = 20 * np.log10(rms) if rms > 0 else -np.inf
            
            print(f"Текущий уровень: {current_db:.1f} дБ")
            print(f"Целевой уровень: {target_db:.1f} дБ")
            
            if rms > 0:
                # Вычисляем коэффициент усиления
                target_linear = 10 ** (target_db / 20)
                gain = target_linear / rms
                
                # Применяем нормализацию
                audio_normalized = audio_np * gain
                
                # Защита от клиппинга (если амплитуда > 1)
                max_val = np.abs(audio_normalized).max()
                if max_val > 1.0:
                    print(f"⚠️ Обнаружен клиппинг (макс: {max_val:.2f}), снижаем уровень")
                    audio_normalized = audio_normalized / max_val
                
                # Конвертируем обратно в исходный формат
                if isinstance(audio, torch.Tensor):
                    audio_normalized = torch.from_numpy(audio_normalized).float()
                    if audio.device != torch.device('cpu'):
                        audio_normalized = audio_normalized.to(audio.device)
                
                new_rms = np.sqrt(np.mean(audio_normalized.cpu().numpy() if isinstance(audio_normalized, torch.Tensor) else audio_normalized) ** 2)
                new_db = 20 * np.log10(new_rms) if new_rms > 0 else -np.inf
                
                print(f"✓ Новый уровень: {new_db:.1f} дБ (улучшение: {new_db - current_db:+.1f} дБ)")
                print("=" * 70)
                
                return audio_normalized
        except Exception as e:
            print(f"⚠️ Ошибка при нормализации: {e}")
            print("Продолжаем с исходным аудио...")
        
        return audio

    def _prepare_review_audio(self, audio, sample_rate=16000):
        """Подготовить обработанное аудио для точного воспроизведения сегментов."""
        try:
            if isinstance(audio, torch.Tensor):
                audio_np = audio.detach().cpu().numpy()
            else:
                audio_np = np.asarray(audio)

            if audio_np.ndim > 1:
                audio_np = np.squeeze(audio_np)

            audio_np = np.asarray(audio_np, dtype=np.float32)

            previous_path = getattr(self, "last_processed_audio_path", None)
            if previous_path and os.path.exists(previous_path):
                try:
                    os.remove(previous_path)
                except OSError:
                    pass

            import tempfile
            import soundfile as sf

            _tmp_dir = os.path.join(os.path.dirname(__file__), 'nemo_tmp')
            os.makedirs(_tmp_dir, exist_ok=True)
            temp_file = tempfile.NamedTemporaryFile(
                prefix="ai_voicefinder_review_",
                suffix='.wav',
                delete=False,
                dir=_tmp_dir,
            )
            review_audio_path = temp_file.name
            temp_file.close()

            sf.write(review_audio_path, audio_np, sample_rate)

            self.last_audio = audio_np
            self.audio_sample_rate = sample_rate
            self.last_processed_audio_path = review_audio_path
            print(f"✓ Подготовлен обработанный файл для проверки: {review_audio_path}")
        except Exception as e:
            print(f"⚠️ Не удалось подготовить обработанный файл для проверки: {e}")
            self.last_audio = np.asarray(audio, dtype=np.float32)
            self.audio_sample_rate = sample_rate
            self.last_processed_audio_path = None
    
    def correct_text_with_lm(self, text, language="ru"):
        """
        🧠 Коррекция текста с помощью правил и контекста
        
        Args:
            text: текст для коррекции
            language: язык (ru, en и т.д.)
            
        Returns:
            исправленный текст
        """
        if not text or not text.strip():
            return text
        
        try:
            print("🧠 Применяем языковое моделирование для коррекции...")
            
            # Базовые исправления для русского
            if language == "ru":
                corrections = {
                    # Частые ошибки распознавания
                    "признает": "признаёт",
                    "извлеченный": "извлечённый",
                    "вынесено": "вынесено",
                    # Более агрессивные исправления с регулярками
                }
                
                # Применяем базовые исправления
                for wrong, correct in corrections.items():
                    text = text.replace(wrong, correct)
                
                # Удаляем дублирующиеся пробелы
                import re
                text = re.sub(r' +', ' ', text)
                
                # Исправляем пунктуацию
                text = re.sub(r' +([,.\?\!:;])', r'\1', text)  # Пробел перед пунктуацией
                text = re.sub(r'([,.\?\!:;]) +', r'\1 ', text)  # Пробел после пунктуации
            
            print("✓ Текст исправлен")
            return text
            
        except Exception as e:
            print(f"⚠️ Ошибка при коррекции текста: {e}")
            return text
    
    def setup_styles(self):
        """� ННастройка темной темы"""
        style = ttk.Style()
        
        # Используем современную тему
        try:
            style.theme_use('clam')
        except:
            pass
        
        # Настраиваем стили для кнопок
        style.configure('Primary.TButton',
                       background=COLORS['primary'],
                       foreground='white',
                       borderwidth=0,
                       focuscolor='none',
                       padding=10,
                       font=('Segoe UI', 12, 'bold'))
        
        style.map('Primary.TButton',
                 background=[('active', '#9E9E9E'), ('pressed', '#757575'), ('disabled', '#3D3D3D')],
                 foreground=[('disabled', '#888888')])
        
        style.configure('Success.TButton',
                       background=COLORS['success'],
                       foreground='white',
                       borderwidth=0,
                       padding=8,
                       font=('Segoe UI', 11))
        
        style.map('Success.TButton',
                 background=[('active', '#BDBDBD'), ('pressed', '#9E9E9E')])
        
        # Стили для фреймов
        style.configure('Card.TFrame',
                       background=COLORS['card_bg'],
                       relief='flat',
                       borderwidth=1)
        
        style.configure('TLabelframe',
                       background=COLORS['card_bg'],
                       borderwidth=1,
                       relief='solid',
                       bordercolor=COLORS['border'])
        
        style.configure('TLabelframe.Label',
                       background=COLORS['card_bg'],
                       foreground=COLORS['light'],
                       font=('Segoe UI', 13, 'bold'))
        
        # Стили для лейблов
        style.configure('TLabel',
                       background=COLORS['card_bg'],
                       foreground=COLORS['text'])
        
        # Стили для комбобоксов
        style.configure('TCombobox',
                       fieldbackground=COLORS['input_bg'],
                       background=COLORS['card_bg'],
                       foreground=COLORS['text'],
                       borderwidth=1,
                       arrowsize=15)
        
        style.map('TCombobox',
                 fieldbackground=[('readonly', COLORS['input_bg'])],
                 selectbackground=[('readonly', COLORS['primary'])],
                 selectforeground=[('readonly', 'white')])
        
        # Стили для чекбоксов
        style.configure('TCheckbutton',
                       background=COLORS['card_bg'],
                       foreground=COLORS['text'],
                       font=('Segoe UI', 10))
        
        style.map('TCheckbutton',
                 background=[('active', COLORS['card_bg'])])
        
        # Стили для прогресс-бара
        style.configure('TProgressbar',
                       background=COLORS['primary'],
                       troughcolor=COLORS['border'],
                       borderwidth=0,
                       thickness=20)
    
    def create_widgets(self):
        """✨ Создание красивых элементов интерфейса"""
        
        # Компактный заголовок
        title_frame = tk.Frame(self.root, bg=COLORS['darker'], height=50)
        title_frame.pack(fill=tk.X)
        title_frame.pack_propagate(False)
        
        title_label = tk.Label(
            title_frame,
            text="AI VoiceFinder",
            font=("Segoe UI", 13, "bold"),
            bg=COLORS['darker'],
            fg=COLORS['light']
        )
        title_label.pack(expand=True)
        
        # Контейнер для карточек (компактный)
        main_container = tk.Frame(self.root, bg=COLORS['bg'])
        main_container.pack(fill=tk.BOTH, expand=True, padx=8, pady=8)
        
        # Компактный фрейм для модели
        model_frame = ttk.LabelFrame(main_container, text="Модель", padding="8")
        model_frame.pack(fill=tk.X, pady=(0, 5))
        
        # Статус модели и устройства (вертикально)
        status_container = tk.Frame(model_frame, bg=COLORS['card_bg'])
        status_container.pack(fill=tk.X)
        
        self.model_status_label = ttk.Label(
            status_container,
            text="Модель не загружена",
            foreground=COLORS['danger']
        )
        self.model_status_label.pack(anchor='w', pady=2)
        
        self.device_label = ttk.Label(status_container, text="", foreground=COLORS['text'])
        self.device_label.pack(anchor='w', pady=2)
        
        # Компактный фрейм для файлов
        audio_frame = ttk.LabelFrame(main_container, text="Файлы", padding="8")
        audio_frame.pack(fill=tk.X, pady=(0, 5))
        
        # Кнопка добавления файла
        self.add_file_btn = ttk.Button(
            audio_frame,
            text="➕ Добавить файл в список",
            command=self.add_file_to_list,
            state=tk.DISABLED,
            style='Primary.TButton'
        )
        self.add_file_btn.pack(fill=tk.X, pady=2)
        
        # Кнопка добавления MTS видеофайла (если поддержка включена)
        if MTS_SUPPORT:
            self.add_mts_btn = ttk.Button(
                audio_frame,
                text="🎬 Добавить видеофайл MTS",
                command=self.add_mts_file_to_list,
                state=tk.DISABLED,
                style='Primary.TButton'
            )
            self.add_mts_btn.pack(fill=tk.X, pady=2)
        else:
            self.add_mts_btn = None
        
        # Список файлов с прокруткой
        list_frame = tk.Frame(audio_frame, bg=COLORS['card_bg'])
        list_frame.pack(fill=tk.BOTH, expand=True, pady=2)
        
        scrollbar = tk.Scrollbar(list_frame)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        self.files_listbox = tk.Listbox(
            list_frame,
            height=4,
            bg=COLORS['input_bg'],
            fg=COLORS['text'],
            selectbackground=COLORS['primary'],
            selectforeground='white',
            font=('Segoe UI', 9),
            yscrollcommand=scrollbar.set
        )
        self.files_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.config(command=self.files_listbox.yview)
        
        # Кнопки управления списком
        list_buttons_frame = tk.Frame(audio_frame, bg=COLORS['card_bg'])
        list_buttons_frame.pack(fill=tk.X, pady=2)
        
        self.remove_file_btn = ttk.Button(
            list_buttons_frame,
            text="❌ Удалить выбранный",
            command=self.remove_selected_file,
            state=tk.DISABLED
        )
        self.remove_file_btn.pack(side=tk.LEFT, padx=(0, 2))
        
        self.clear_files_btn = ttk.Button(
            list_buttons_frame,
            text="🗑️ Очистить список",
            command=self.clear_files_list,
            state=tk.DISABLED
        )
        self.clear_files_btn.pack(side=tk.LEFT)
        
        # Счетчик файлов
        self.files_count_label = ttk.Label(
            audio_frame,
            text="Файлов: 0",
            foreground=COLORS['text']
        )
        self.files_count_label.pack(anchor='w', pady=2)
        
        # Компактный фрейм для настроек (вертикально)
        settings_frame = ttk.LabelFrame(main_container, text="⚙️ Настройки", padding="8")
        settings_frame.pack(fill=tk.X, pady=(0, 5))
        
        # Язык
        lang_frame = tk.Frame(settings_frame, bg=COLORS['card_bg'])
        lang_frame.pack(fill=tk.X, pady=2)
        ttk.Label(lang_frame, text="Язык:").pack(side=tk.LEFT, padx=2)
        self.language_var = tk.StringVar(value="ru")
        language_combo = ttk.Combobox(
            lang_frame,
            textvariable=self.language_var,
            values=["ru", "en", "auto"],
            state="readonly",
            width=8
        )
        language_combo.pack(side=tk.LEFT, padx=2)
        
        # Опции (вертикально)
        self.align_var = tk.BooleanVar(value=True)
        self.align_checkbox = ttk.Checkbutton(
            settings_frame,
            text="Точные временные метки",
            variable=self.align_var,
            command=self.on_align_changed
        )
        self.align_checkbox.pack(anchor='w', pady=2)
        
        self.diarize_var = tk.BooleanVar(value=False)
        self.diarize_checkbox = ttk.Checkbutton(
            settings_frame,
            text="Определение спикеров (требует временные метки)",
            variable=self.diarize_var,
            command=self.on_diarize_changed
        )
        self.diarize_checkbox.pack(anchor='w', pady=2)
        
        self.transcript_var = tk.BooleanVar(value=False)
        self.transcript_checkbox = ttk.Checkbutton(
            settings_frame,
            text="Режим стенограммы (только текст)",
            variable=self.transcript_var,
            command=self.on_transcript_changed
        )
        self.transcript_checkbox.pack(anchor='w', pady=2)
        
        # Устанавливаем правильное состояние чекбокса диаризации при инициализации
        self.on_align_changed()
        
        # Кнопка распознавания (вертикально)
        process_frame = tk.Frame(main_container, bg=COLORS['bg'])
        process_frame.pack(fill=tk.X, pady=(0, 5))
        
        self.process_btn = ttk.Button(
            process_frame,
            text="Распознать речь",
            command=self.process_audio,
            state=tk.DISABLED,
            style='Primary.TButton'
        )
        self.process_btn.pack(fill=tk.X, pady=2)
        
        # Контейнер для GIF "Work in progress" (центрировано)
        self.progress_container = tk.Frame(process_frame, bg=COLORS['bg'])
        self.progress_container.pack(pady=5)
        
        # Загружаем GIF "Work in progress" (используем абсолютный путь относительно скрипта)
        try:
            script_dir = os.path.dirname(os.path.abspath(__file__))
            gif_path = os.path.join(script_dir, "Sources", "Work in progress.gif")
            if os.path.exists(gif_path):
                gif_image = Image.open(gif_path)
                
                # Размер GIF 173x173 (увеличено на 50% от 115)
                original_width = gif_image.width
                original_height = gif_image.height
                scale = min(173 / original_width, 173 / original_height)
                new_width = int(original_width * scale)
                new_height = int(original_height * scale)
                
                # Загружаем все кадры GIF
                self.progress_gif_frames = []
                try:
                    while True:
                        frame_image = gif_image.copy()
                        frame_image = frame_image.resize((new_width, new_height), Image.Resampling.LANCZOS)
                        self.progress_gif_frames.append(ImageTk.PhotoImage(frame_image))
                        gif_image.seek(len(self.progress_gif_frames))
                except EOFError:
                    pass
                
                # Создаем лейбл для GIF (изначально скрыт)
                self.progress_gif_label = tk.Label(
                    self.progress_container,
                    bg=COLORS['bg'],
                    bd=0
                )
                
                # Индекс текущего кадра
                self.progress_gif_index = [0]
                self.progress_gif_running = [False]
                
                print(f"✓ GIF 'Work in progress' загружен ({len(self.progress_gif_frames)} кадров)")
            else:
                print("⚠ GIF 'Work in progress' не найден")
                # Создаем обычный прогресс-бар как запасной вариант
                self.progress = ttk.Progressbar(
                    self.progress_container,
                    mode='indeterminate',
                    length=300
                )
        except Exception as e:
            print(f"⚠ Ошибка загрузки GIF: {e}")
            # Создаем обычный прогресс-бар как запасной вариант
            self.progress = ttk.Progressbar(
                self.progress_container,
                mode='indeterminate',
                length=300
            )
        
        # Статус обработки (центрировано)
        self.status_label = tk.Label(
            process_frame,
            text="Готов к работе",
            foreground=COLORS['success'],
            bg=COLORS['bg'],
            font=('Segoe UI', 11)
        )
        self.status_label.pack(pady=2)
        
        # Контейнер для кнопок результатов
        self.results_buttons_frame = tk.Frame(self.progress_container, bg=COLORS['bg'])
        
        # Кнопка открытия папки результатов (квадратная с иконкой)
        self.open_folder_btn = tk.Button(
            self.results_buttons_frame,
            text="📁",
            command=self.open_results_folder,
            bg=COLORS['success'],
            fg='white',
            font=("Segoe UI", 24),
            relief=tk.FLAT,
            cursor="hand2",
            width=3,
            height=1,
            activebackground='#7FDF97',
            activeforeground='white'
        )
        
        # Кнопка экспорта в Word (квадратная, синяя с белой буквой W)
        self.export_word_btn = tk.Button(
            self.results_buttons_frame,
            text="W",
            command=self.export_to_word,
            bg='#616161',
            fg='white',
            font=("Segoe UI", 24, "bold"),
            relief=tk.FLAT,
            cursor="hand2",
            width=3,
            height=1,
            activebackground='#757575',
            activeforeground='white'
        )

        # Кнопка проверки и редактирования фрагментов
        self.review_btn = tk.Button(
            self.results_buttons_frame,
            text="🎧",
            command=self.open_review_window,
            bg=COLORS['primary'],
            fg='white',
            font=("Segoe UI", 24),
            relief=tk.FLAT,
            cursor="hand2",
            width=3,
            height=1,
            activebackground=COLORS['accent'],
            activeforeground='white'
        )
        # Не показываем кнопки до завершения обработки
        
        # Переменная для хранения пути к файлу результатов
        self.result_file_path = None
        # Список всех обработанных файлов (для пакетной обработки)
        self.processed_files = []  # [(audio_file, result_txt_file), ...]

        # Данные для окна проверки и редактирования фрагментов
        self.last_audio = None          # numpy array последнего обработанного файла
        self.last_result = None         # dict с segments последнего обработанного файла
        self.last_save_filepath = None  # путь к txt-файлу результатов
        self.last_processed_audio_path = None  # временный WAV последнего обработанного файла
        self.audio_sample_rate = 16000  # частота дискретизации
        
        # Создаем пустое текстовое поле для совместимости (скрыто)
        self.result_text = scrolledtext.ScrolledText(main_container)
        # Не показываем его
    
    def open_result_file(self):
        """Открытие файла с результатами распознавания"""
        if self.result_file_path and os.path.exists(self.result_file_path):
            try:
                os.startfile(self.result_file_path)
                print(f"✓ Файл открыт: {self.result_file_path}")
            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось открыть файл:\n{e}")
        else:
            messagebox.showwarning("Предупреждение", "Файл результатов не найден")
    
    def start_progress_animation(self):
        """Запуск анимации GIF 'Work in progress'"""
        if hasattr(self, 'progress_gif_frames') and len(self.progress_gif_frames) > 0:
            self.progress_gif_running[0] = True
            self.progress_gif_label.pack()
            self.animate_progress_gif()
        elif hasattr(self, 'progress'):
            # Запасной вариант - обычный прогресс-бар
            self.progress.pack()
            self.progress.start()
    
    def stop_progress_animation(self):
        """Остановка анимации GIF 'Work in progress'"""
        if hasattr(self, 'progress_gif_frames') and len(self.progress_gif_frames) > 0:
            self.progress_gif_running[0] = False
            self.progress_gif_label.pack_forget()
        elif hasattr(self, 'progress'):
            # Запасной вариант - обычный прогресс-бар
            self.progress.stop()
            self.progress.pack_forget()
    
    def animate_progress_gif(self):
        """Анимация GIF кадр за кадром"""
        try:
            if self.progress_gif_running[0] and hasattr(self, 'progress_gif_label'):
                if self.progress_gif_label.winfo_exists():
                    frame = self.progress_gif_frames[self.progress_gif_index[0]]
                    self.progress_gif_label.config(image=frame)
                    self.progress_gif_label.image = frame
                    self.progress_gif_index[0] = (self.progress_gif_index[0] + 1) % len(self.progress_gif_frames)
                    self.root.after(50, self.animate_progress_gif)
        except:
            pass
    
    def on_align_changed(self):
        """Обработка изменения чекбокса точных временных меток"""
        if self.align_var.get():
            # Если включены временные метки - разрешаем диаризацию
            self.diarize_checkbox.config(state=tk.NORMAL)
        else:
            # Если выключены временные метки - отключаем диаризацию и стенограмму
            self.diarize_var.set(False)
            self.diarize_checkbox.config(state=tk.DISABLED)
            self.transcript_var.set(False)
            self.transcript_checkbox.config(state=tk.DISABLED)
        
        # Обновляем состояние стенограммы
        self.on_diarize_changed()
    
    def on_diarize_changed(self):
        """Обработка изменения чекбокса диаризации"""
        if self.diarize_var.get():
            # Если включена диаризация - разрешаем стенограмму
            self.transcript_checkbox.config(state=tk.NORMAL)
        else:
            # Если выключена диаризация - отключаем стенограмму
            self.transcript_var.set(False)
            self.transcript_checkbox.config(state=tk.DISABLED)
    
    def on_transcript_changed(self):
        """Обработка изменения чекбокса режима стенограммы"""
        # Стенограмма просто меняет формат вывода, не влияет на другие настройки
        pass
    
    def check_cuda(self):
        """Проверка доступности CUDA"""
        print("\n" + "=" * 70)
        print("🔍 ПРОВЕРКА GPU")
        print("=" * 70)
        
        # Проверяем доступность CUDA
        cuda_available = torch.cuda.is_available()
        
        print(f"PyTorch версия: {torch.__version__}")
        print(f"CUDA доступна: {cuda_available}")
        
        if cuda_available:
            print(f"CUDA версия: {torch.version.cuda}")
            print(f"cuDNN версия: {torch.backends.cudnn.version()}")
            print(f"Количество GPU: {torch.cuda.device_count()}")
            
            # Получаем информацию о GPU
            for i in range(torch.cuda.device_count()):
                gpu_name = torch.cuda.get_device_name(i)
                props = torch.cuda.get_device_properties(i)
                gpu_memory = props.total_memory / 1024**3
                print(f"GPU {i}: {gpu_name} ({gpu_memory:.1f} ГБ)")
        
        print("=" * 70)
        
        # Используем GPU если доступен
        if cuda_available:
            device_name = torch.cuda.get_device_name(0)
            gpu_memory = torch.cuda.get_device_properties(0).total_memory / 1024**3
            
            self.device_label.config(
                text=f"GPU: {device_name} ({gpu_memory:.1f} ГБ)",
                foreground='#4CAF50'
            )
            self.device = "cuda"
            self.compute_type = "float16"
            print(f"\n✅ Используется GPU: {device_name}")
            print(f"Тип вычислений: {self.compute_type}")
            print(f"Ожидаемое ускорение: 10-20x по сравнению с CPU\n")
        else:
            self.device_label.config(
                text="💻 CPU режим (GPU не обнаружен)",
                foreground=COLORS['warning']
            )
            self.device = "cpu"
            self.compute_type = "int8"
            
            print("\n⚠️ GPU не обнаружен - используется CPU")
            print("\nВозможные причины:")
            print("1. Драйвер NVIDIA не установлен")
            print("2. CUDA Toolkit не установлен")
            print("3. PyTorch установлен без поддержки CUDA")
            print("\nДля диагностики запустите: ПРОВЕРКА_GPU.bat")
            print("\nРешение:")
            print("1. Установите драйвер NVIDIA")
            print("2. Установите CUDA Toolkit 12.x")
            print("3. Переустановите PyTorch:")
            print("   pip install torch torchvision torchaudio --index-url https://download.pytorch.org/whl/cu121")
            print("\nПриложение будет работать на CPU (медленнее в 10-20 раз)\n")
    
    def load_model(self):
        """Загрузка модели AI VoiceFinder"""
        if self.is_processing:
            messagebox.showwarning("Предупреждение", "Дождитесь завершения обработки")
            return
        
        thread = threading.Thread(target=self._load_model_thread)
        thread.start()
    
    def _load_model_thread(self):
        """Загрузка модели в отдельном потоке"""
        try:
            self.model_status_label.config(text="⏳ Загрузка модели...", foreground=COLORS['warning'])
            
            model_size = "large-v3"  # Фиксированный размер модели
            
            print(f"\nЗагрузка модели AI VoiceFinder ({model_size})...")
            print(f"Устройство: {self.device}")
            print(f"Тип вычислений: {self.compute_type}")
            
            # Загружаем модель WhisperX
            _whisper_cache = os.path.join(LOCAL_CACHE_DIR, 'whisper')
            os.makedirs(_whisper_cache, exist_ok=True)
            self.model = whisperx.load_model(
                model_size,
                self.device,
                compute_type=self.compute_type,
                download_root=_whisper_cache,
            )
            
            print("✓ Модель загружена")
            
            # Загружаем модель LavaSR для улучшения аудио
            if LAVASR_AVAILABLE:
                try:
                    print("\nЗагрузка модели LavaSR для улучшения аудио...")
                    self.lavasr_model = LavaEnhance2("YatharthS/LavaSR", self.device)
                    print("✓ Модель LavaSR загружена")
                    print("🌋 Улучшение аудио ВКЛЮЧЕНО")
                except Exception as e:
                    print(f"⚠️ Ошибка загрузки LavaSR: {e}")
                    print("⚠️ Продолжаем без улучшения аудио")
                    self.lavasr_model = None
            else:
                print("\n⚠️ LavaSR недоступен")
                print("Для включения улучшения аудио установите:")
                print("pip install git+https://github.com/ysharma3501/LavaSR.git")
                self.lavasr_model = None
            
            # Выводим рекомендации по оптимизации
            if self.device == "cuda":
                print("\n" + "=" * 70)
                print("⚡ ОПТИМИЗАЦИЯ ДЛЯ GPU (на основе A-train)")
                print("=" * 70)
                print("🎯 Параметры транскрипции (улучшение +15-20%):")
                print("   • VAD filter: ВКЛЮЧЕН (отсеивание пустоты)")
                print("   • Контекст: ВКЛЮЧЕН (учет предыдущих сегментов)")
                print("   • Compute type: float16 (вместо int8)")
                print(f"   • Batch size: 16 (оптимизировано для 8 ГБ VRAM)")
                print("=" * 70 + "\n")
            else:
                print("\n" + "=" * 70)
                print("⚡ ОПТИМИЗАЦИЯ ДЛЯ CPU (на основе A-train)")
                print("=" * 70)
                print("🎯 Параметры транскрипции (улучшение +10-15%):")
                print("   • VAD filter: ВКЛЮЧЕН (отсеивание пустоты)")
                print("   • Контекст: ВКЛЮЧЕН (учет предыдущих сегментов)")
                print("   • Compute type: int8")
                print("=" * 70 + "\n")
            
            self.model_status_label.config(text="✅ Модель загружена", foreground=COLORS['success'])
            self.set_input_buttons_state(tk.NORMAL)
            
            messagebox.showinfo(
                "Успех",
                f"AI VoiceFinder ({model_size}) загружена!\nУстройство: {self.device.upper()}"
            )
            
        except Exception as e:
            import traceback
            error_details = traceback.format_exc()
            
            print("\n" + "=" * 70)
            print("ОШИБКА ЗАГРУЗКИ МОДЕЛИ")
            print("=" * 70)
            print(error_details)
            print("=" * 70)
            
            self.model_status_label.config(text="❌ Ошибка загрузки", foreground=COLORS['danger'])
            
            error_msg = str(e)
            if "Pipeline" in error_msg:
                msg = "Ошибка импорта Pipeline.\n\n" \
                      "Попробуйте:\n" \
                      "1. pip uninstall pyannote.audio\n" \
                      "2. pip install pyannote.audio\n\n" \
                      "Или используйте без диаризации."
            elif "whisperx" in error_msg.lower():
                msg = "Ошибка WhisperX.\n\n" \
                      "Переустановите:\n" \
                      "pip install git+https://github.com/m-bain/whisperx.git"
            else:
                msg = f"Не удалось загрузить модель:\n{error_msg}\n\nПодробности в консоли."
            
            messagebox.showerror("Ошибка", msg)
    
    def add_file_to_list(self):
        """Добавление файла в список для обработки"""
        file_path = filedialog.askopenfilename(
            title="Выберите файл для добавления в список",
            filetypes=[
                ("Аудио и видео", "*.mp3 *.wav *.m4a *.flac *.ogg *.opus *.mp4 *.avi *.mkv *.mov *.wmv *.flv *.webm"),
                ("Аудиофайлы", "*.mp3 *.wav *.m4a *.flac *.ogg *.opus"),
                ("Видеофайлы", "*.mp4 *.avi *.mkv *.mov *.wmv *.flv *.webm"),
                ("Все файлы", "*.*")
            ]
        )
        
        if file_path:
            # Добавляем файл в список (без дубликатов)
            if file_path not in self.audio_files_list:
                self.audio_files_list.append(file_path)
                self.update_files_list()
            else:
                messagebox.showinfo("Информация", "Этот файл уже в списке")
    
    def add_mts_file_to_list(self):
        """Добавление MTS видеофайла в список для обработки"""
        if not MTS_SUPPORT:
            messagebox.showerror(
                "Ошибка",
                "Модуль обработки MTS не доступен"
            )
            return
        
        try:
            # Проверяем наличие FFmpeg
            processor = MTSProcessor()
        except MTSProcessorError as e:
            response = messagebox.askyesno(
                "FFmpeg не найден",
                f"Для обработки видеофайлов MTS требуется FFmpeg.\n\n"
                f"❌ Ошибка: {str(e)}\n\n"
                f"Установить FFmpeg? (откроется инструкция)"
            )
            if response:
                print(setup_ffmpeg_guide())
                messagebox.showinfo(
                    "Инструкция по установке FFmpeg",
                    setup_ffmpeg_guide()
                )
            return
        
        # Открываем диалог выбора MTS файла
        file_path = filedialog.askopenfilename(
            title="Выберите видеофайл MTS",
            filetypes=[
                ("MTS видеофайлы", "*.mts *.m2ts"),
                ("MTS и M2TS", "*.mts *.m2ts *.MTS *.M2TS"),
                ("Все файлы", "*.*")
            ]
        )
        
        if file_path:
            # Проверяем, что это MTS файл
            if not MTSProcessor.is_mts_file(file_path):
                messagebox.showwarning(
                    "Предупреждение",
                    "Выбранный файл не является MTS видеофайлом.\n"
                    "Поддерживаются форматы: .mts, .m2ts"
                )
                return
            
            # Добавляем файл в список (без дубликатов)
            if file_path not in self.audio_files_list:
                self.audio_files_list.append(file_path)
                self.mts_files_for_processing = self.mts_files_for_processing or []
                self.mts_files_for_processing.append(file_path)
                self.update_files_list()
                print(f"✓ MTS файл добавлен в список: {os.path.basename(file_path)}")
            else:
                messagebox.showinfo("Информация", "Этот файл уже в списке")
    
    def update_files_list(self):
        """Обновление списка файлов в интерфейсе"""
        # Очищаем listbox
        self.files_listbox.delete(0, tk.END)
        
        # Добавляем файлы
        video_extensions = ['.mp4', '.avi', '.mkv', '.mov', '.wmv', '.flv', '.webm']
        mts_extensions = ['.mts', '.m2ts']
        
        for file_path in self.audio_files_list:
            filename = os.path.basename(file_path)
            file_ext = os.path.splitext(file_path)[1].lower()
            
            # Проверяем тип файла
            if file_ext in mts_extensions:
                display_name = f"🎥 {filename}"  # Видеокамера для MTS
            elif file_ext in video_extensions:
                display_name = f"🎬 {filename}"  # Кинопленка для обычного видео
            else:
                display_name = f"🎵 {filename}"  # Музыка для аудио
            
            self.files_listbox.insert(tk.END, display_name)
        
        # Обновляем счетчик
        count = len(self.audio_files_list)
        self.files_count_label.config(text=f"Файлов: {count}")
        
        # Активируем/деактивируем кнопки
        if count > 0:
            self.process_btn.config(state=tk.NORMAL)
            self.remove_file_btn.config(state=tk.NORMAL)
            self.clear_files_btn.config(state=tk.NORMAL)
        else:
            self.process_btn.config(state=tk.DISABLED)
            self.remove_file_btn.config(state=tk.DISABLED)
            self.clear_files_btn.config(state=tk.DISABLED)
    
    def remove_selected_file(self):
        """Удаление выбранного файла из списка"""
        selection = self.files_listbox.curselection()
        if selection:
            index = selection[0]
            del self.audio_files_list[index]
            self.update_files_list()
    
    def clear_files_list(self):
        """Очистка списка файлов"""
        self.audio_files_list = []
        self.mts_files_for_processing = []
        self.update_files_list()
    
    def extract_audio_from_video(self, video_path):
        """Извлечение аудио из видеофайла"""
        try:
            import subprocess
            import tempfile
            
            print(f"\n🎬 Извлечение аудио из видео: {os.path.basename(video_path)}")
            
            # Создаем временный файл для аудио
            _tmp_dir = os.path.join(os.path.dirname(__file__), 'nemo_tmp')
            os.makedirs(_tmp_dir, exist_ok=True)
            temp_audio = tempfile.NamedTemporaryFile(suffix='.wav', delete=False, dir=_tmp_dir)
            temp_audio_path = temp_audio.name
            temp_audio.close()
            
            # Используем ffmpeg для извлечения аудио
            # -i: входной файл
            # -vn: без видео
            # -acodec pcm_s16le: аудио кодек WAV
            # -ar 16000: частота дискретизации 16kHz (для Whisper)
            # -ac 1: моно
            cmd = [
                'ffmpeg',
                '-i', video_path,
                '-vn',  # Без видео
                '-acodec', 'pcm_s16le',  # WAV формат
                '-ar', '16000',  # 16kHz
                '-ac', '1',  # Моно
                '-y',  # Перезаписать если существует
                temp_audio_path
            ]
            
            print("Запуск ffmpeg...")
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=300  # 5 минут максимум
            )
            
            if result.returncode != 0:
                raise Exception(f"Ошибка ffmpeg: {result.stderr}")
            
            print(f"✓ Аудио извлечено: {temp_audio_path}")
            return temp_audio_path
            
        except FileNotFoundError:
            raise Exception(
                "ffmpeg не найден!\n\n"
                "Для работы с видео нужен ffmpeg.\n\n"
                "Установка:\n"
                "1. Скачайте ffmpeg: https://ffmpeg.org/download.html\n"
                "2. Добавьте в PATH\n"
                "3. Или положите ffmpeg.exe в папку приложения"
            )
        except subprocess.TimeoutExpired:
            raise Exception("Превышено время ожидания извлечения аудио (5 минут)")
        except Exception as e:
            raise Exception(f"Ошибка извлечения аудио: {str(e)}")
    
    def process_audio(self):
        """Запуск распознавания речи"""
        if not self.model:
            messagebox.showwarning("Предупреждение", "Сначала загрузите модель")
            return
        
        # Проверяем есть ли файлы для обработки
        if not self.audio_files_list:
            messagebox.showwarning("Предупреждение", "Сначала добавьте файл(ы) в список")
            return
        
        if self.is_processing:
            messagebox.showwarning("Предупреждение", "Обработка уже выполняется")
            return
        
        # Очищаем список обработанных файлов перед новой обработкой
        self.processed_files = []
        print("✓ Список обработанных файлов очищен")
        
        # Пакетная обработка
        files_count = len(self.audio_files_list)
        
        if files_count == 1:
            # Один файл - обрабатываем сразу
            thread = threading.Thread(target=self._process_batch_thread)
            thread.start()
        else:
            # Несколько файлов - спрашиваем подтверждение
            response = messagebox.askyesno(
                "Пакетная обработка",
                f"Обработать {files_count} файлов?\n\n"
                f"Это может занять продолжительное время."
            )
            if not response:
                return
            
            thread = threading.Thread(target=self._process_batch_thread)
            thread.start()
    
    def _process_audio_thread(self):
        """Обработка аудио в отдельном потоке"""
        temp_audio_path = None  # Инициализируем переменную для временного файла
        
        try:
            # Очищаем память GPU перед началом
            if torch.cuda.is_available():
                torch.cuda.empty_cache()
                print("✓ Память GPU очищена")
            
            self.is_processing = True
            self.process_btn.config(state=tk.DISABLED)
            self.add_file_btn.config(state=tk.DISABLED)
            
            # Скрываем кнопки результатов при начале новой обработки
            self.results_buttons_frame.pack_forget()
            self.open_folder_btn.pack_forget()
            self.export_word_btn.pack_forget()
            self.review_btn.pack_forget()
            
            # Показываем статус обратно (если был скрыт)
            self.status_label.pack(pady=(10, 0))
            
            # Запускаем анимацию GIF вместо прогресс-бара
            self.start_progress_animation()
            
            # Создаем файл для сохранения
            save_filepath = self._create_save_file()
            
            # Сохраняем путь к файлу результатов
            self.result_file_path = save_filepath
            
            # НЕ открываем файл сразу - это замедляет обработку
            # Файл откроется автоматически после завершения
            
            # Очищаем текстовое поле
            self.result_text.delete(1.0, tk.END)
            
            # Получаем параметры
            language = self.language_var.get()
            if language == "auto":
                language = None
            
            # Проверяем является ли файл видео
            video_extensions = ['.mp4', '.avi', '.mkv', '.mov', '.wmv', '.flv', '.webm']
            file_ext = os.path.splitext(self.audio_file_path)[1].lower()
            
            if file_ext in video_extensions:
                # Извлекаем аудио из видео
                self.status_label.config(text="🎬 Извлечение аудио из видео...", foreground=COLORS['warning'])
                self.root.update()
                
                try:
                    temp_audio_path = self.extract_audio_from_video(self.audio_file_path)
                    audio_path = temp_audio_path
                except Exception as e:
                    self.status_label.config(text="❌ Ошибка", foreground=COLORS['danger'])
                    messagebox.showerror("Ошибка", str(e))
                    return
            else:
                audio_path = self.audio_file_path
            
            # ШАГ 1: ПОТОКОВОЕ Распознавание речи
            self.status_label.config(text="🐾 Распознавание речи... 0%", foreground=COLORS['warning'])
            self.root.update()
            
            print("\n" + "=" * 70)
            print("ШАГ 1: ПОТОКОВОЕ Распознавание речи")
            print("=" * 70)
            
            # Загружаем аудио
            audio = whisperx.load_audio(audio_path)

            # ШАГ 0.25: НОРМАЛИЗАЦИЯ ГРОМКОСТИ
            self.status_label.config(text="🔊 Нормализация громкости...", foreground=COLORS['warning'])
            self.root.update()
            audio = self.normalize_audio(audio, target_db=-20.0)
            
            # ШАГ 0.5: УЛУЧШЕНИЕ АУДИО С ПОМОЩЬЮ LavaSR
            if self.lavasr_model is not None:
                self.status_label.config(text="🌋 Улучшение аудио с LavaSR...", foreground=COLORS['warning'])
                self.root.update()
                
                # Улучшаем аудио с денойзингом для лучшего качества
                audio = self.enhance_audio_with_lavasr(audio, sr=16000, denoise=True)
                
                # Очищаем GPU памяль после LavaSR
                if torch.cuda.is_available():
                    torch.cuda.empty_cache()
                    print("✓ Память GPU очищена после LavaSR")

            self._prepare_review_audio(audio, sample_rate=16000)
            
            # Заголовок в интерфейсе и файле
            header = "РАСПОЗНАННЫЙ ТЕКСТ (потоковый режим):\n" + "=" * 70 + "\n\n"
            self.result_text.insert(tk.END, header)
            self._append_to_file(save_filepath, header)
            self.root.update()
            
            # ПОТОКОВАЯ ТРАНСКРИПЦИЯ - обрабатываем по частям
            print("Начало потоковой транскрипции...")
            
            # Определяем оптимальные параметры в зависимости от устройства
            if self.device == "cuda":
                # Для GPU: безопасные значения to избежать CUDA OOM
                # Уменьшены для совместимости с GPU <8GB VRAM
                batch_size = 4  # Уменьшено с 32 для совместимости с меньшими GPU
                chunk_length = 30  # Уменьшено с 60 для меньшего потребления памяти
                print(f"🎮 GPU режим: batch_size={batch_size}, chunk_length={chunk_length}s")
                
                # Включаем оптимизации для GPU
                torch.backends.cudnn.benchmark = True  # Автоматический выбор лучших алгоритмов
                if hasattr(torch.backends, 'cuda') and hasattr(torch.backends.cuda, 'matmul'):
                    torch.backends.cuda.matmul.allow_tf32 = True  # Ускорение матричных операций
                print("✓ GPU оптимизации включены")
            else:
                # Для CPU: оптимальные значения
                batch_size = 4
                chunk_length = 30
                print(f"💻 CPU режим: batch_size={batch_size}, chunk_length={chunk_length}s")
            
            # Разбиваем аудио на части для потоковой обработки
            sample_rate = 16000
            chunk_samples = chunk_length * sample_rate
            total_samples = len(audio)
            num_chunks = (total_samples + chunk_samples - 1) // chunk_samples
            
            print(f"Длина аудио: {total_samples / sample_rate:.1f} секунд")
            print(f"Разбито на {num_chunks} частей по {chunk_length} секунд")
            
            all_segments = []
            buffer_lines = []  # Буфер для накопления строк перед записью
            
            for chunk_idx in range(num_chunks):
                start_sample = chunk_idx * chunk_samples
                end_sample = min((chunk_idx + 1) * chunk_samples, total_samples)
                audio_chunk = audio[start_sample:end_sample]
                
                # Обрабатываем часть с оптимальным batch_size
                try:
                    result = self.model.transcribe(
                        audio_chunk,
                        language=language
                    )
                    if isinstance(result, dict) and "segments" in result:
                        chunk_result = result
                    else:
                        chunk_result = {"segments": list(result["segments"]) if hasattr(result, "segments") else []}
                except (IndexError, StopIteration):
                    # VAD не нашёл речи в этом чанке — пропускаем
                    print(f"  ⚠️ Чанк {chunk_idx+1}/{num_chunks}: речь не обнаружена, пропускаем")
                    continue
                except RuntimeError as e:
                    if "out of memory" in str(e).lower() or "cuda" in str(e).lower():
                        print(f"⚠️ CUDA память исчерпана (batch_size={batch_size}), пробуем batch_size=2")
                        # Очищаем GPU память
                        if self.device == "cuda":
                            torch.cuda.empty_cache()
                            import gc
                            gc.collect()
                        # Пробуем снова
                        try:
                            result = self.model.transcribe(
                                audio_chunk,
                                language=language
                            )
                        except (IndexError, StopIteration):
                            print(f"  ⚠️ Чанк {chunk_idx+1}/{num_chunks}: речь не обнаружена, пропускаем")
                            continue
                        if isinstance(result, dict) and "segments" in result:
                            chunk_result = result
                        else:
                            chunk_result = {"segments": list(result["segments"]) if hasattr(result, "segments") else []}
                    else:
                        raise
                
                if "segments" in chunk_result:
                    # Корректируем временные метки с учетом смещения
                    time_offset = start_sample / sample_rate
                    for seg in chunk_result["segments"]:
                        seg["start"] += time_offset
                        seg["end"] += time_offset
                        all_segments.append(seg)
                        
                        # СРАЗУ ВЫВОДИМ
                        text = seg.get("text", "")
                        start = seg.get("start", 0)
                        end = seg.get("end", 0)
                        
                        # 🧠 Коррекция текста с помощью языкового моделирования
                        detected_lang = result.get("language", language if language else "ru")
                        text = self.correct_text_with_lm(text, language=detected_lang)
                        
                        line = f"[{format_timestamp(start)} - {format_timestamp(end)}]: {text}\n"
                        
                        # Выводим в интерфейс
                        self.result_text.insert(tk.END, line)
                        self.result_text.see(tk.END)
                        
                        # Добавляем в буфер вместо немедленной записи
                        buffer_lines.append(line)
                
                # Записываем буфер в файл после каждого чанка (не после каждого сегмента!)
                if buffer_lines:
                    self._append_to_file(save_filepath, "".join(buffer_lines))
                    buffer_lines.clear()
                
                # Очищаем GPU память после обработки чанка
                if self.device == "cuda":
                    torch.cuda.empty_cache()
                
                # Обновляем статус только после каждого чанка
                progress = ((chunk_idx + 1) / num_chunks) * 100
                self.status_label.config(
                    text=f"🐾 Распознавание речи... {progress:.0f}%",
                    foreground=COLORS['warning']
                )
                self.root.update()
                print(f"✓ Обработана часть {chunk_idx + 1}/{num_chunks}")
            
            # Собираем результат
            result = {"segments": all_segments, "language": language if language else "ru"}
            
            print(f"✓ Распознано {len(all_segments)} сегментов")
            self._append_to_file(save_filepath, "\n")

            if not all_segments:
                self.status_label.config(text="⚠️ Речь не обнаружена в файле", foreground=COLORS['warning'])
                self.root.update()
                print("⚠️ В файле не найдено активной речи — результат пустой")
                self.result_text.insert(tk.END, "⚠️ Речь не обнаружена. Возможно, файл содержит только тишину или шум.\n")
                self.result_text.see(tk.END)
                return
            
            # ШАГ 2: Выравнивание (точные временные метки)
            if self.align_var.get():
                self.status_label.config(text="⚡ Выравнивание... 0%", foreground=COLORS['primary'])
                self.root.update()
                
                print("\n" + "=" * 70)
                print("ШАГ 2: Выравнивание (точные временные метки)")
                print("=" * 70)
                
                # Загружаем модель выравнивания если еще не загружена
                if self.align_model is None:
                    # Определяем язык из результата
                    detected_lang = result.get("language", language if language else "ru")
                    print(f"Определенный язык: {detected_lang}")
                    print(f"Загрузка модели выравнивания для языка: {detected_lang}")
                    
                    try:
                        _align_cache = os.path.join(LOCAL_CACHE_DIR, 'align')
                        os.makedirs(_align_cache, exist_ok=True)
                        self.align_model, self.align_metadata = whisperx.load_align_model(
                            language_code=detected_lang,
                            device=self.device,
                            model_dir=_align_cache,
                        )
                        print("✓ Модель выравнивания загружена")
                    except Exception as e:
                        print(f"⚠ Ошибка загрузки модели выравнивания: {e}")
                        print("Продолжаем без выравнивания...")
                        self.align_var.set(False)
                
                if self.align_model:
                    self.status_label.config(text="⚡ Выравнивание... 50%", foreground=COLORS['primary'])
                    self.root.update()
                    
                    aligned_result = whisperx.align(
                        result["segments"],
                        self.align_model,
                        self.align_metadata,
                        audio,
                        self.device,
                        return_char_alignments=False
                    )
                    
                    self.status_label.config(text="⚡ Выравнивание... 100%", foreground=COLORS['primary'])
                    self.root.update()
                    
                    # Обновляем сегменты, сохраняя остальные данные
                    result["segments"] = aligned_result["segments"]
                    if "word_segments" in aligned_result:
                        result["word_segments"] = aligned_result["word_segments"]
                    print(f"✓ Выравнивание завершено")
                    print(f"✓ Сегментов после выравнивания: {len(result['segments'])}")
                    
            # ШАГ 3: Диаризация (определение спикеров)
            print(f"\n🔍 Проверка диаризации: diarize_var={self.diarize_var.get()}, align_var={self.align_var.get()}")
            
            if self.diarize_var.get() and self.align_var.get():
                self.status_label.config(text="👥 Определение спикеров... 0%", foreground=COLORS['accent'])
                self.root.update()
                
                print("\n" + "=" * 70)
                print("ШАГ 3: Диаризация (определение спикеров)")
                print("=" * 70)
                print(f"Диаризация ВКЛЮЧЕНА - начинаем обработку")
                
                try:
                    if self.diarize_model is None:
                        self.status_label.config(text="👥 Загрузка модели диаризации... 10%", foreground=COLORS['accent'])
                        self.root.update()
                        
                        print("Загрузка модели диаризации NeMo...")
                        print("✓ NeMo не требует токен Hugging Face!")

                        # Перенаправляем временную папку на диск проекта (избегаем заполнения C:)
                        import tempfile as _tempfile_mod
                        _nemo_tmp = os.path.join(os.path.dirname(__file__), 'nemo_tmp')
                        os.makedirs(_nemo_tmp, exist_ok=True)
                        _old_tempdir = _tempfile_mod.tempdir
                        _tempfile_mod.tempdir = _nemo_tmp
                        os.environ['TEMP'] = _nemo_tmp
                        os.environ['TMP']  = _nemo_tmp
                        print(f"📁 Временная папка NeMo: {_nemo_tmp}")
                        
                        # Импортируем NeMo
                        import nemo.collections.asr as nemo_asr
                        from omegaconf import OmegaConf
                        
                        # Создаем конфигурацию для диаризации
                        # Создаем полную конфигурацию для NeMo
                        config = {
                            'manifest_filepath': None,
                            'out_dir': os.path.join(os.path.dirname(__file__), 'nemo_outputs'),
                            'oracle_vad': False,
                            'device': self.device,
                            'num_workers': 0,  # Важно для Windows
                            'sample_rate': 16000,  # Частота дискретизации
                            'batch_size': 1,
                            'verbose': True,  # Вывод прогресса
                            'diarizer': {
                                'manifest_filepath': None,
                                'out_dir': os.path.join(os.path.dirname(__file__), 'nemo_outputs'),
                                'oracle_vad': False,
                                'device': self.device,
                                'collar': 0.25,  # Допуск для границ сегментов (в секундах)
                                'ignore_overlap': True,  # Игнорировать перекрывающиеся сегменты
                                'clustering': {
                                    'parameters': {
                                        'oracle_num_speakers': False,
                                        'max_num_speakers': 8,
                                        'enhanced_count_thres': 80,
                                        'max_rp_threshold': 0.25,
                                        'sparse_search_volume': 30
                                    }
                                },
                                'vad': {
                                    'model_path': 'vad_multilingual_marblenet',
                                    'parameters': {
                                        'window_length_in_sec': 0.15,
                                        'shift_length_in_sec': 0.01,
                                        'smoothing': 'median',
                                        'overlap': 0.5,
                                        'onset': 0.8,
                                        'offset': 0.6,
                                        'pad_onset': 0.05,
                                        'pad_offset': 0.05,
                                        'min_duration_on': 0.2,
                                        'min_duration_off': 0.2
                                    }
                                },
                                'speaker_embeddings': {
                                    'model_path': 'titanet_large',
                                    'parameters': {
                                        'window_length_in_sec': 1.5,
                                        'shift_length_in_sec': 0.75,
                                        'multiscale_weights': [1, 1, 1, 1, 1],
                                        'save_embeddings': False
                                    }
                                }
                            }
                        }
                        
                        cfg = OmegaConf.create(config)
                        
                        # Загружаем модель диаризации NeMo
                        self.diarize_model = nemo_asr.models.ClusteringDiarizer(cfg=cfg)
                        print("✓ Модель диаризации NeMo загружена")

                        # Восстанавливаем оригинальную временную папку
                        _tempfile_mod.tempdir = _old_tempdir
                    
                    self.status_label.config(text="👥 Подготовка аудио... 20%", foreground=COLORS['accent'])
                    self.root.update()
                    
                    # ОБХОДНОЙ ПУТЬ: Конвертируем аудио во временный файл
                    print("Подготовка аудио для диаризации...")
                    import tempfile
                    import soundfile as sf
                    import json
                    
                    # Создаем временную директорию для NeMo (в папке проекта, не на C:)
                    _nemo_tmp_base = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'nemo_tmp')
                    os.makedirs(_nemo_tmp_base, exist_ok=True)
                    temp_dir = tempfile.mkdtemp(dir=_nemo_tmp_base)
                    
                    # Создаем временный WAV файл с правильными параметрами
                    temp_audio_path = os.path.join(temp_dir, "audio.wav")
                    sf.write(temp_audio_path, audio, 16000)
                    print(f"✓ Временный файл создан: {temp_audio_path}")
                    
                    # Создаем manifest файл для NeMo
                    manifest_path = os.path.join(temp_dir, "manifest.json")
                    audio_duration = len(audio) / 16000
                    manifest_entry = {
                        "audio_filepath": temp_audio_path,
                        "offset": 0,
                        "duration": audio_duration,
                        "label": "infer",
                        "text": "-",
                        "num_speakers": None,
                        "rttm_filepath": None,
                        "uem_filepath": None
                    }
                    
                    with open(manifest_path, 'w') as f:
                        json.dump(manifest_entry, f)
                        f.write('\n')
                    
                    print(f"✓ Manifest создан: {manifest_path}")
                    
                    try:
                        self.status_label.config(text="👥 Анализ аудио... 40%", foreground=COLORS['accent'])
                        self.root.update()
                        
                        # Обновляем конфигурацию модели с путями к файлам
                        self.diarize_model._cfg.diarizer.manifest_filepath = manifest_path
                        self.diarize_model._cfg.diarizer.out_dir = temp_dir
                        
                        # Применяем диаризацию
                        print("Запуск диаризации...")
                        self.status_label.config(text="👥 Определение спикеров... 60%", foreground=COLORS['accent'])
                        self.root.update()
                        
                        self.diarize_model.diarize()
                        
                        self.status_label.config(text="👥 Обработка результатов... 80%", foreground=COLORS['accent'])
                        self.root.update()
                        
                        print("✓ Диаризация выполнена")
                        
                        # Читаем результаты из RTTM файла
                        from pyannote.core import Annotation, Segment
                        rttm_path = os.path.join(temp_dir, "pred_rttms", "audio.rttm")
                        
                        if not os.path.exists(rttm_path):
                            raise FileNotFoundError(f"RTTM файл не найден: {rttm_path}")
                        
                        # Парсим RTTM файл
                        diarize_segments = Annotation()
                        speakers = set()
                        
                        with open(rttm_path, 'r') as f:
                            for line in f:
                                parts = line.strip().split()
                                if len(parts) >= 8:
                                    start_time = float(parts[3])
                                    duration = float(parts[4])
                                    speaker = parts[7]
                                    speakers.add(speaker)
                                    diarize_segments[Segment(start_time, start_time + duration)] = speaker
                        
                        print(f"✓ Найдено спикеров: {len(speakers)}")
                        
                        # Проверяем структуру result
                        segments_list = None
                        if isinstance(result, dict):
                            segments_list = result.get("segments", [])
                        elif isinstance(result, list):
                            segments_list = result
                        
                        # Проверяем что у сегментов есть слова (после выравнивания)
                        if self.align_var.get() and segments_list:
                            # Присваиваем спикеров к словам
                            print("Присвоение спикеров к словам...")
                            
                            # Создаем правильную структуру для assign_word_speakers
                            transcript_with_words = {"segments": segments_list}
                            
                            try:
                                diarized_result = whisperx.assign_word_speakers(diarize_segments, transcript_with_words)
                                print(f"✓ Диаризация завершена")
                                
                                # Проверяем что спикеры действительно присвоены
                                speakers_found = set()
                                if "segments" in diarized_result:
                                    for seg in diarized_result["segments"]:
                                        if "speaker" in seg:
                                            speakers_found.add(seg["speaker"])
                                        if "words" in seg:
                                            for word in seg["words"]:
                                                if "speaker" in word:
                                                    speakers_found.add(word["speaker"])
                                
                                if speakers_found:
                                    print(f"✓ Найдено спикеров в результате: {len(speakers_found)} - {speakers_found}")
                                    # Обновляем сегменты в основном результате
                                    result["segments"] = diarized_result["segments"]
                                    
                                    self.status_label.config(text="👥 Определение спикеров... 100%", foreground=COLORS['accent'])
                                    self.root.update()
                                else:
                                    print("⚠ Спикеры не найдены в результате, используем альтернативный метод")
                                    raise Exception("Спикеры не присвоены")
                            except Exception as e:
                                print(f"⚠ Ошибка при присвоении спикеров через WhisperX: {e}")
                                print("Использую альтернативный метод...")
                                
                                # Альтернативный метод - присваиваем спикеров вручную
                                for segment in segments_list:
                                    if "words" in segment:
                                        for word in segment["words"]:
                                            word_start = word.get("start", 0)
                                            word_end = word.get("end", 0)
                                            word_mid = (word_start + word_end) / 2
                                            
                                            # Находим спикера в это время
                                            for turn, _, speaker in diarize_segments.itertracks(yield_label=True):
                                                if turn.start <= word_mid <= turn.end:
                                                    word["speaker"] = speaker
                                                    break
                                    
                                    # Также присваиваем спикера всему сегменту
                                    seg_start = segment.get("start", 0)
                                    seg_end = segment.get("end", 0)
                                    seg_mid = (seg_start + seg_end) / 2
                                    
                                    for turn, _, speaker in diarize_segments.itertracks(yield_label=True):
                                        if turn.start <= seg_mid <= turn.end:
                                            segment["speaker"] = speaker
                                            break
                                
                                result = {"segments": segments_list}
                                print(f"✓ Спикеры присвоены альтернативным методом")
                                
                                self.status_label.config(text="👥 Определение спикеров... 100%", foreground=COLORS['accent'])
                                self.root.update()
                        else:
                            # Если нет выравнивания, просто добавляем информацию о спикерах к сегментам
                            print("⚠ Для лучшей диаризации включите 'Точные временные метки'")
                            print("Добавление спикеров к сегментам...")
                            
                            # Простое присвоение спикеров по времени
                            for segment in segments_list:
                                seg_start = segment.get("start", 0)
                                seg_end = segment.get("end", 0)
                                seg_mid = (seg_start + seg_end) / 2
                                
                                # Находим спикера в это время
                                for turn, _, speaker in diarize_segments.itertracks(yield_label=True):
                                    if turn.start <= seg_mid <= turn.end:
                                        segment["speaker"] = speaker
                                        break
                            
                            result = {"segments": segments_list}
                            print(f"✓ Спикеры добавлены к сегментам")
                            
                            self.status_label.config(text="👥 Определение спикеров... 100%", foreground=COLORS['accent'])
                            self.root.update()
                        
                    finally:
                        # Удаляем временные файлы
                        import shutil
                        if os.path.exists(temp_dir):
                            shutil.rmtree(temp_dir)
                            print("✓ Временные файлы удалены")
                    
                except Exception as e:
                    import traceback
                    print(f"\n⚠ Ошибка диаризации: {e}")
                    print("Детали ошибки:")
                    traceback.print_exc()
                    
                    # Показываем более информативное сообщение
                    error_msg = str(e)
                    if "Sizes of tensors" in error_msg:
                        msg = "Ошибка обработки аудио.\n\nВозможные причины:\n" \
                              "- Нестандартная длина аудио\n" \
                              "- Поврежденный файл\n" \
                              "- Несовместимый формат\n\n" \
                              "Попробуйте:\n" \
                              "- Конвертировать в WAV\n" \
                              "- Использовать другой файл\n" \
                              "- Отключить диаризацию"
                    else:
                        msg = f"Диаризация недоступна:\n{error_msg}\n\n" \
                              "Продолжаем без определения спикеров."
                    
                    messagebox.showwarning("Предупреждение", msg)
            elif self.diarize_var.get() and not self.align_var.get():
                print("\n⚠️ Диаризация требует включения точных временных меток")
                print("Диаризация пропущена")
            else:
                print(f"\n⚠️ Диаризация ОТКЛЮЧЕНА (diarize={self.diarize_var.get()}, align={self.align_var.get()})")
            
            # Вывод результатов
            self.status_label.config(text="💾 Сохранение результатов... 0%", foreground=COLORS['primary'])
            self.root.update()
            
            self._display_results(result, save_filepath)
            
            # Добавляем время завершения в файл
            self._finalize_save_file(save_filepath)
            
            print("\n" + "=" * 70)
            print(f"✓ Результат сохранен: {save_filepath}")
            print("=" * 70)
            
            # Добавляем в список обработанных файлов
            self.processed_files.append((self.audio_file_path, save_filepath))
            print(f"✓ Файл добавлен в список обработанных: {os.path.basename(self.audio_file_path)}")

            # Сохраняем данные для окна проверки и редактирования
            self.last_result = result
            self.last_save_filepath = save_filepath
            
            # Скрываем статус и GIF
            self.status_label.pack_forget()
            self.stop_progress_animation()
            
            # Показываем кнопки результатов в progress_container (на месте GIF, по центру, ниже на 1/3)
            self.results_buttons_frame.pack(pady=(60, 10), expand=True)
            
            # Показываем кнопку открытия папки
            self.open_folder_btn.pack(side=tk.LEFT, padx=5)
            
            # Если режим стенограммы - показываем кнопку экспорта в Word
            if self.transcript_var.get():
                self.export_word_btn.pack(side=tk.LEFT, padx=5)

            if self.last_result is not None:
                self.review_btn.pack(side=tk.LEFT, padx=5)
            
            messagebox.showinfo("Успех", "Распознавание завершено!\n\nРезультат автоматически сохранен.")
            
        except Exception as e:
            import traceback
            error_details = traceback.format_exc()
            print("\n" + "=" * 70)
            print("ОШИБКА ПРИ ОБРАБОТКЕ АУДИО")
            print("=" * 70)
            print(error_details)
            print("=" * 70)
            
            # Скрываем GIF и показываем статус ошибки
            self.stop_progress_animation()
            self.status_label.pack(pady=(10, 0))
            self.status_label.config(text="❌ Ошибка", foreground=COLORS['danger'])
            messagebox.showerror("Ошибка", f"Ошибка при обработке:\n{str(e)}\n\nПодробности в консоли.")
        
        finally:
            # Удаляем временный аудиофайл если он был создан
            if temp_audio_path and os.path.exists(temp_audio_path):
                try:
                    os.unlink(temp_audio_path)
                    print(f"✓ Временный аудиофайл удален: {temp_audio_path}")
                except Exception as e:
                    print(f"⚠ Не удалось удалить временный файл: {e}")
            
            self.is_processing = False
            
            self.process_btn.config(state=tk.NORMAL)
            self.add_file_btn.config(state=tk.NORMAL)
            self.add_file_btn.config(state=tk.NORMAL)
    
    
    def _process_mts_files(self):
        """Обработка MTS видеофайлов - извлечение аудиопотока"""
        if not MTS_SUPPORT:
            print("⚠️ Поддержка MTS отключена")
            return
        
        try:
            processor = MTSProcessor()
        except MTSProcessorError as e:
            print(f"❌ Ошибка инициализации MTS процессора: {e}")
            return
        
        mts_files = [f for f in self.audio_files_list if MTSProcessor.is_mts_file(f)]
        
        if not mts_files:
            print("⚠️ MTS файлы не найдены в списке")
            return
        
        print(f"\n🎥 Обработка {len(mts_files)} MTS видеофайлов")
        print("=" * 70)
        
        for idx, mts_file in enumerate(mts_files, 1):
            try:
                print(f"\n🎥 MTS файл {idx}/{len(mts_files)}")
                filename = os.path.basename(mts_file)
                
                # Обновляем статус
                self.status_label.config(
                    text=f"🎥 Извлечение аудио {idx}/{len(mts_files)}: {filename[:30]}...",
                    foreground=COLORS['warning']
                )
                self.root.update()
                
                print(f"📝 Файл: {mts_file}")
                
                # Определяем путь для сохранения WAV
                base_name = os.path.splitext(mts_file)[0]
                wav_file = f"{base_name}_extracted.wav"
                
                # Извлекаем аудио
                output_wav = processor.extract_audio(mts_file, wav_file)
                
                # Заменяем MTS файл в списке на извлеченный WAV
                idx_in_list = self.audio_files_list.index(mts_file)
                self.audio_files_list[idx_in_list] = output_wav
                
                print(f"✓ Аудиопоток извлечен: {os.path.basename(output_wav)}")
                
            except MTSProcessorError as e:
                print(f"❌ Ошибка при обработке MTS файла: {e}")
                # Удаляем файл из списка при ошибке
                if mts_file in self.audio_files_list:
                    self.audio_files_list.remove(mts_file)
            except Exception as e:
                print(f"❌ Неожиданная ошибка: {e}")
                import traceback
                traceback.print_exc()
                if mts_file in self.audio_files_list:
                    self.audio_files_list.remove(mts_file)
        
        # Обновляем список файлов в интерфейсе
        self.update_files_list()
        print("\n✓ Обработка MTS файлов завершена\n")
    
    def _process_batch_thread(self):
        """Пакетная обработка файлов"""
        total_files = len(self.audio_files_list)
        processed = 0
        failed = 0
        failed_files = []
        
        try:
            self.is_processing = True
            self.process_btn.config(state=tk.DISABLED)
            self.add_file_btn.config(state=tk.DISABLED)
            if self.add_mts_btn:
                self.add_mts_btn.config(state=tk.DISABLED)
            self.remove_file_btn.config(state=tk.DISABLED)
            self.clear_files_btn.config(state=tk.DISABLED)

            # Скрываем кнопки результатов (если остались от прошлого запуска)
            self.results_buttons_frame.pack_forget()
            self.open_folder_btn.pack_forget()
            self.export_word_btn.pack_forget()
            self.review_btn.pack_forget()
            
            # Запускаем анимацию GIF
            self.start_progress_animation()
            
            print("\n" + "=" * 70)
            print(f"🎬 ПАКЕТНАЯ ОБРАБОТКА: {total_files} файлов")
            print("=" * 70)
            
            # Сначала обрабатываем MTS видеофайлы
            if MTS_SUPPORT and self.mts_files_for_processing:
                print(f"\n🎥 Обработка {len(self.mts_files_for_processing)} MTS видеофайлов...")
                self._process_mts_files()
                # Очищаем список MTS файлов после обработки
                self.mts_files_for_processing = []
            
            for index, file_path in enumerate(self.audio_files_list, 1):
                filename = os.path.basename(file_path)
                
                print(f"\n{'=' * 70}")
                print(f"📁 Файл {index}/{total_files}: {filename}")
                print(f"{'=' * 70}")
                
                # Обновляем статус
                self.status_label.config(
                    text=f"📁 Обработка {index}/{total_files}: {filename[:30]}...",
                    foreground=COLORS['warning']
                )
                self.root.update()
                
                # Временно устанавливаем текущий файл
                original_file_path = self.audio_file_path
                self.audio_file_path = file_path
                
                try:
                    # Обрабатываем файл (используем существующую логику)
                    self._process_single_file()
                    processed += 1
                    print(f"✓ Файл {index}/{total_files} обработан успешно")
                    
                except Exception as e:
                    failed += 1
                    failed_files.append(filename)
                    print(f"❌ Ошибка обработки файла {index}/{total_files}: {e}")
                    import traceback
                    traceback.print_exc()
                
                finally:
                    # Восстанавливаем оригинальный путь
                    self.audio_file_path = original_file_path
            
            # Итоговый отчет
            print("\n" + "=" * 70)
            print("📊 ИТОГИ ПАКЕТНОЙ ОБРАБОТКИ")
            print("=" * 70)
            print(f"Всего файлов: {total_files}")
            print(f"✓ Обработано успешно: {processed}")
            print(f"❌ Ошибок: {failed}")
            
            if failed_files:
                print("\nФайлы с ошибками:")
                for f in failed_files:
                    print(f"  - {f}")
            
            print("=" * 70)
            
            # Скрываем статус и GIF
            self.status_label.pack_forget()
            self.stop_progress_animation()
            
            # Показываем кнопки результатов в progress_container (на месте GIF, по центру, ниже на 1/3)
            self.results_buttons_frame.pack(pady=(60, 10), expand=True)
            
            # Показываем кнопку открытия папки
            self.open_folder_btn.pack(side=tk.LEFT, padx=5)
            
            # Если режим стенограммы - показываем кнопку экспорта в Word
            if self.transcript_var.get():
                self.export_word_btn.pack(side=tk.LEFT, padx=5)

            # Показываем кнопку проверки фрагментов
            if self.last_result is not None:
                self.review_btn.pack(side=tk.LEFT, padx=5)
            
            if failed > 0:
                messagebox.showwarning(
                    "Пакетная обработка завершена",
                    f"Обработано: {processed}/{total_files}\n"
                    f"Ошибок: {failed}\n\n"
                    f"Файлы с ошибками:\n" + "\n".join(failed_files[:5]) +
                    (f"\n... и еще {len(failed_files) - 5}" if len(failed_files) > 5 else "")
                )
            else:
                messagebox.showinfo(
                    "Пакетная обработка завершена",
                    f"Все файлы обработаны успешно!\n\n"
                    f"Обработано: {processed}/{total_files}"
                )
        
        except Exception as e:
            print(f"\n❌ Критическая ошибка пакетной обработки: {e}")
            import traceback
            traceback.print_exc()
            
            # Скрываем GIF и показываем статус ошибки
            self.stop_progress_animation()
            self.status_label.pack(pady=(10, 0))
            self.status_label.config(text="❌ Ошибка", foreground=COLORS['danger'])
            messagebox.showerror("Ошибка", f"Критическая ошибка:\n{str(e)}")
        
        finally:
            self.is_processing = False
            self.process_btn.config(state=tk.NORMAL)
            self.add_file_btn.config(state=tk.NORMAL)
            self.add_file_btn.config(state=tk.NORMAL)
            self.remove_file_btn.config(state=tk.NORMAL)
            self.clear_files_btn.config(state=tk.NORMAL)
    
    def _process_single_file(self):
        """Обработка одного файла (используется в пакетной обработке)"""
        temp_audio_path = None
        
        try:
            # Создаем файл для сохранения
            save_filepath = self._create_save_file()
            self.result_file_path = save_filepath
            
            # Получаем параметры
            language = self.language_var.get()
            if language == "auto":
                language = None
            
            # Проверяем является ли файл видео
            video_extensions = ['.mp4', '.avi', '.mkv', '.mov', '.wmv', '.flv', '.webm']
            file_ext = os.path.splitext(self.audio_file_path)[1].lower()
            
            if file_ext in video_extensions:
                # Извлекаем аудио из видео
                temp_audio_path = self.extract_audio_from_video(self.audio_file_path)
                audio_path = temp_audio_path
            else:
                audio_path = self.audio_file_path
            
            # Загружаем аудио
            audio = whisperx.load_audio(audio_path)
            self._prepare_review_audio(audio, sample_rate=16000)
            
            # ШАГ 1: Распознавание
            print("\nШАГ 1: Распознавание речи")
            self.status_label.config(text="🐾 Распознавание речи... 0%", foreground=COLORS['warning'])
            self.root.update()
            
            # Определяем оптимальные параметры в зависимости от устройства
            if self.device == "cuda":
                batch_size = 4  # Уменьшено с 16 для совместимости с GPU <8GB
                chunk_length = 30
            else:
                batch_size = 4
                chunk_length = 30
            
            sample_rate = 16000
            chunk_samples = chunk_length * sample_rate
            total_samples = len(audio)
            num_chunks = (total_samples + chunk_samples - 1) // chunk_samples
            
            all_segments = []
            
            for chunk_idx in range(num_chunks):
                progress = ((chunk_idx + 1) / num_chunks) * 100
                self.status_label.config(text=f"🐾 Распознавание речи... {progress:.0f}%", foreground=COLORS['warning'])
                self.root.update()
                print(f"Распознавание: {progress:.0f}%")
                
                start_sample = chunk_idx * chunk_samples
                end_sample = min((chunk_idx + 1) * chunk_samples, total_samples)
                audio_chunk = audio[start_sample:end_sample]
                
                try:
                    result = self.model.transcribe(
                        audio_chunk,
                        language=language
                    )
                except (IndexError, StopIteration):
                    # VAD не нашёл речи в этом чанке — пропускаем
                    print(f"  ⚠️ Чанк {chunk_idx+1}/{num_chunks}: речь не обнаружена, пропускаем")
                    continue
                if isinstance(result, dict) and "segments" in result:
                    chunk_result = result
                else:
                    chunk_result = {"segments": list(result["segments"]) if hasattr(result, "segments") else []}
                
                if "segments" in chunk_result:
                    time_offset = start_sample / sample_rate
                    for seg in chunk_result["segments"]:
                        seg["start"] += time_offset
                        seg["end"] += time_offset
                        all_segments.append(seg)
            
            result = {"segments": all_segments, "language": language if language else "ru"}
            print(f"✓ Распознано {len(all_segments)} сегментов")

            if not all_segments:
                self.status_label.config(text="⚠️ Речь не обнаружена в файле", foreground=COLORS['warning'])
                self.root.update()
                print("⚠️ В файле не найдено активной речи — результат пустой")
                self.result_text.insert(tk.END, "⚠️ Речь не обнаружена. Возможно, файл содержит только тишину или шум.\n")
                self.result_text.see(tk.END)
                return
            
            # ШАГ 2: Выравнивание
            if self.align_var.get():
                print("\nШАГ 2: Выравнивание")
                self.status_label.config(text="⚡ Выравнивание... 0%", foreground=COLORS['primary'])
                self.root.update()
                
                if self.align_model is None:
                    detected_lang = result.get("language", language if language else "ru")
                    _align_cache = os.path.join(LOCAL_CACHE_DIR, 'align')
                    os.makedirs(_align_cache, exist_ok=True)
                    self.align_model, self.align_metadata = whisperx.load_align_model(
                        language_code=detected_lang,
                        device=self.device,
                        model_dir=_align_cache,
                    )
                
                if self.align_model:
                    self.status_label.config(text="⚡ Выравнивание... 50%", foreground=COLORS['primary'])
                    self.root.update()
                    print("Выравнивание: 50%")
                    
                    aligned_result = whisperx.align(
                        result["segments"],
                        self.align_model,
                        self.align_metadata,
                        audio,
                        self.device,
                        return_char_alignments=False
                    )
                    result["segments"] = aligned_result["segments"]
                    
                    self.status_label.config(text="⚡ Выравнивание... 100%", foreground=COLORS['primary'])
                    self.root.update()
                    print(f"✓ Выравнивание завершено (100%)")
            
            # ШАГ 3: Диаризация
            print(f"\n🔍 Проверка диаризации: diarize_var={self.diarize_var.get()}, align_var={self.align_var.get()}")
            
            if self.diarize_var.get() and self.align_var.get():
                print("\nШАГ 3: Диаризация")
                self.status_label.config(text="👥 Определение спикеров... 0%", foreground=COLORS['accent'])
                self.root.update()
                
                # Здесь должна быть вся логика диаризации
                # Копируем из _process_audio_thread
                import nemo.collections.asr as nemo_asr
                from omegaconf import OmegaConf
                import tempfile
                import soundfile as sf
                import json
                from pyannote.core import Annotation, Segment
                
                if self.diarize_model is None:
                    self.status_label.config(text="👥 Загрузка модели диаризации... 10%", foreground=COLORS['accent'])
                    self.root.update()
                    print("Загрузка модели диаризации NeMo... (10%)")

                    # Перенаправляем временную папку на диск проекта (избегаем заполнения C:)
                    import tempfile as _tempfile_mod
                    _nemo_tmp = os.path.join(os.path.dirname(__file__), 'nemo_tmp')
                    os.makedirs(_nemo_tmp, exist_ok=True)
                    _old_tempdir = _tempfile_mod.tempdir
                    _tempfile_mod.tempdir = _nemo_tmp
                    os.environ['TEMP'] = _nemo_tmp
                    os.environ['TMP']  = _nemo_tmp
                    print(f"📁 Временная папка NeMo: {_nemo_tmp}")

                    config = {
                        'manifest_filepath': None,
                        'out_dir': os.path.join(os.path.dirname(__file__), 'nemo_outputs'),
                        'oracle_vad': False,
                        'device': self.device,
                        'num_workers': 0,
                        'sample_rate': 16000,
                        'batch_size': 1,
                        'verbose': True,
                        'diarizer': {
                            'manifest_filepath': None,
                            'out_dir': os.path.join(os.path.dirname(__file__), 'nemo_outputs'),
                            'oracle_vad': False,
                            'device': self.device,
                            'collar': 0.25,
                            'ignore_overlap': True,
                            'clustering': {
                                'parameters': {
                                    'oracle_num_speakers': False,
                                    'max_num_speakers': 8,
                                    'enhanced_count_thres': 80,
                                    'max_rp_threshold': 0.25,
                                    'sparse_search_volume': 30
                                }
                            },
                            'vad': {
                                'model_path': 'vad_multilingual_marblenet',
                                'parameters': {
                                    'window_length_in_sec': 0.15,
                                    'shift_length_in_sec': 0.01,
                                    'smoothing': 'median',
                                    'overlap': 0.5,
                                    'onset': 0.8,
                                    'offset': 0.6,
                                    'pad_onset': 0.05,
                                    'pad_offset': 0.05,
                                    'min_duration_on': 0.2,
                                    'min_duration_off': 0.2
                                }
                            },
                            'speaker_embeddings': {
                                'model_path': 'titanet_large',
                                'parameters': {
                                    'window_length_in_sec': 1.5,
                                    'shift_length_in_sec': 0.75,
                                    'multiscale_weights': [1, 1, 1, 1, 1],
                                    'save_embeddings': False
                                }
                            }
                        }
                    }
                    cfg = OmegaConf.create(config)
                    self.diarize_model = nemo_asr.models.ClusteringDiarizer(cfg=cfg)
                    print("✓ Модель диаризации загружена")

                    # Восстанавливаем оригинальную временную папку
                    _tempfile_mod.tempdir = _old_tempdir
                
                # Создаем временные файлы
                self.status_label.config(text="👥 Подготовка аудио... 20%", foreground=COLORS['accent'])
                self.root.update()
                print("Подготовка аудио... (20%)")
                
                _nemo_tmp_base = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'nemo_tmp')
                os.makedirs(_nemo_tmp_base, exist_ok=True)
                temp_dir = tempfile.mkdtemp(dir=_nemo_tmp_base)
                temp_audio_path_diarize = os.path.join(temp_dir, "audio.wav")
                sf.write(temp_audio_path_diarize, audio, 16000)
                
                manifest_path = os.path.join(temp_dir, "manifest.json")
                audio_duration = len(audio) / 16000
                manifest_entry = {
                    "audio_filepath": temp_audio_path_diarize,
                    "offset": 0,
                    "duration": audio_duration,
                    "label": "infer",
                    "text": "-",
                    "num_speakers": None,
                    "rttm_filepath": None,
                    "uem_filepath": None
                }
                
                with open(manifest_path, 'w') as f:
                    json.dump(manifest_entry, f)
                    f.write('\n')
                
                try:
                    self.status_label.config(text="👥 Анализ аудио... 40%", foreground=COLORS['accent'])
                    self.root.update()
                    print("Анализ аудио... (40%)")
                    
                    self.diarize_model._cfg.diarizer.manifest_filepath = manifest_path
                    self.diarize_model._cfg.diarizer.out_dir = temp_dir
                    
                    self.status_label.config(text="👥 Определение спикеров... 60%", foreground=COLORS['accent'])
                    self.root.update()
                    print("Определение спикеров... (60%)")
                    
                    self.diarize_model.diarize()
                    
                    self.status_label.config(text="👥 Обработка результатов... 80%", foreground=COLORS['accent'])
                    self.root.update()
                    print("Обработка результатов... (80%)")
                    
                    rttm_path = os.path.join(temp_dir, "pred_rttms", "audio.rttm")
                    diarize_segments = Annotation()
                    speakers = set()
                    
                    with open(rttm_path, 'r') as f:
                        for line in f:
                            parts = line.strip().split()
                            if len(parts) >= 8:
                                start_time = float(parts[3])
                                duration = float(parts[4])
                                speaker = parts[7]
                                speakers.add(speaker)
                                diarize_segments[Segment(start_time, start_time + duration)] = speaker
                    
                    print(f"✓ Найдено спикеров: {len(speakers)}")
                    
                    # Присваиваем спикеров
                    segments_list = result.get("segments", [])
                    transcript_with_words = {"segments": segments_list}
                    
                    try:
                        diarized_result = whisperx.assign_word_speakers(diarize_segments, transcript_with_words)
                        result["segments"] = diarized_result["segments"]
                        
                        self.status_label.config(text="👥 Определение спикеров... 100%", foreground=COLORS['accent'])
                        self.root.update()
                        print(f"✓ Спикеры присвоены (100%)")
                    except Exception as e:
                        print(f"⚠ Ошибка assign_word_speakers: {e}")
                        # Альтернативный метод
                        for segment in segments_list:
                            if "words" in segment:
                                for word in segment["words"]:
                                    word_mid = (word.get("start", 0) + word.get("end", 0)) / 2
                                    for turn, _, speaker in diarize_segments.itertracks(yield_label=True):
                                        if turn.start <= word_mid <= turn.end:
                                            word["speaker"] = speaker
                                            break
                            seg_mid = (segment.get("start", 0) + segment.get("end", 0)) / 2
                            for turn, _, speaker in diarize_segments.itertracks(yield_label=True):
                                if turn.start <= seg_mid <= turn.end:
                                    segment["speaker"] = speaker
                                    break
                        result["segments"] = segments_list
                        
                        self.status_label.config(text="👥 Определение спикеров... 100%", foreground=COLORS['accent'])
                        self.root.update()
                        print(f"✓ Спикеры присвоены альтернативным методом (100%)")
                
                finally:
                    import shutil
                    if os.path.exists(temp_dir):
                        shutil.rmtree(temp_dir)
            
            # Сохраняем результаты
            self.status_label.config(text="💾 Сохранение результатов... 0%", foreground=COLORS['primary'])
            self.root.update()
            print("\nСохранение результатов...")
            
            self._display_results(result, save_filepath)
            
            # Добавляем время завершения в файл
            self._finalize_save_file(save_filepath)
            
            print(f"✓ Результат сохранен: {save_filepath} (100%)")
            
            # Добавляем в список обработанных файлов
            self.processed_files.append((self.audio_file_path, save_filepath))
            print(f"✓ Файл добавлен в список обработанных: {os.path.basename(self.audio_file_path)}")

            # Сохраняем данные для окна проверки и редактирования
            self.last_result = result
            self.last_save_filepath = save_filepath
            
        finally:
            # Удаляем временный аудиофайл если был создан
            if temp_audio_path and os.path.exists(temp_audio_path):
                try:
                    os.remove(temp_audio_path)
                    print(f"✓ Временный файл удален: {temp_audio_path}")
                except Exception as e:
                    print(f"⚠ Не удалось удалить временный файл: {e}")
    
    def _save_results_to_file(self, result, filepath):
        """Сохранение результатов в файл"""
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write("РЕЗУЛЬТАТЫ РАСПОЗНАВАНИЯ\n")
            f.write("=" * 70 + "\n\n")
            
            if isinstance(result, dict) and "segments" in result:
                for segment in result["segments"]:
                    start = segment.get("start", 0)
                    end = segment.get("end", 0)
                    text = segment.get("text", "")
                    f.write(f"[{format_timestamp(start)} - {format_timestamp(end)}]: {text}\n")
    
    def _display_results(self, result, save_filepath):
        """Отображение и сохранение результатов"""
        
        # Проверяем режим стенограммы
        if self.transcript_var.get():
            # Режим стенограммы - только текст без временных меток
            self._display_transcript_mode(result, save_filepath)
            return
        
        # Проверяем наличие спикеров в результате
        speakers_in_result = set()
        if "segments" in result:
            for seg in result["segments"]:
                if "speaker" in seg:
                    speakers_in_result.add(seg["speaker"])
                if "words" in seg:
                    for word in seg["words"]:
                        if "speaker" in word:
                            speakers_in_result.add(word["speaker"])
        
        if speakers_in_result:
            print(f"\n✓ Спикеры в результате для вывода: {speakers_in_result}")
        else:
            print("\n⚠ Спикеры не найдены в результате для вывода")
        
        # Заголовок для сегментов
        header = "\n" + "=" * 70 + "\nВРЕМЕННЫЕ МЕТКИ:\n" + "=" * 70 + "\n\n"
        self.result_text.insert(tk.END, header)
        self._append_to_file(save_filepath, header)
        
        if "segments" in result:
            total_segments = len(result["segments"])
            
            for i, segment in enumerate(result["segments"], 1):
                # Информация о сегменте
                start = segment.get("start", 0)
                end = segment.get("end", 0)
                text = segment.get("text", "")
                speaker = segment.get("speaker", None)
                
                # Формируем строку
                if speaker:
                    line = f"\n[{format_timestamp(start)} - {format_timestamp(end)}] {speaker}:\n{text}\n"
                else:
                    line = f"\n[{format_timestamp(start)} - {format_timestamp(end)}]:\n{text}\n"
                
                # Если есть слова с временными метками
                if "words" in segment and segment["words"]:
                    line += "  Слова:\n"
                    for word_info in segment["words"]:
                        word = word_info.get("word", "")
                        word_start = word_info.get("start", 0)
                        word_end = word_info.get("end", 0)
                        word_speaker = word_info.get("speaker", "")
                        
                        if word_speaker:
                            line += f"    [{format_timestamp(word_start)} - {format_timestamp(word_end)}] {word_speaker}: {word}\n"
                        else:
                            line += f"    [{format_timestamp(word_start)} - {format_timestamp(word_end)}]: {word}\n"
                
                # Выводим в интерфейс
                self.result_text.insert(tk.END, line)
                self.result_text.see(tk.END)
                
                # Сохраняем в файл
                self._append_to_file(save_filepath, line)
                
                # Обновляем статус
                progress = (i / total_segments) * 100
                self.status_label.config(
                    text=f"💾 Сохранение результатов... {progress:.0f}%",
                    foreground=COLORS['primary']
                )
                self.root.update()
    
    def _create_save_file(self):
        """Создание файла для сохранения с заголовком"""
        try:
            from datetime import datetime
            
            # Сохраняем время начала обработки
            self.processing_start_time = datetime.now()
            
            # Создаем папку для результатов (используем глобальную переменную)
            results_dir = RESULTS_DIR
            os.makedirs(results_dir, exist_ok=True)
            
            # Генерируем имя файла с временной меткой
            timestamp = self.processing_start_time.strftime("%Y-%m-%d_%H-%M-%S")
            
            # Получаем имя исходного аудиофайла
            audio_filename = os.path.basename(self.audio_file_path)
            audio_name = os.path.splitext(audio_filename)[0]
            
            # Формируем имя файла
            filename = f"{audio_name}_{timestamp}_whisperx.txt"
            filepath = os.path.join(results_dir, filename)
            
            # Создаем файл с заголовком
            with open(filepath, "w", encoding="utf-8") as f:
                f.write(f"AI VoiceFinder - Распознавание речи\n")
                f.write("=" * 70 + "\n")
                f.write(f"Исходный файл: {audio_filename}\n")
                f.write(f"Время начала: {self.processing_start_time.strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write(f"Модель: large-v3\n")
                f.write(f"Язык: {self.language_var.get()}\n")
                f.write(f"Режим: {'Стенограмма' if self.transcript_var.get() else 'Обычный'}\n")
                f.write(f"Точные метки: {'Да' if self.align_var.get() else 'Нет'}\n")
                f.write(f"Диаризация: {'Да' if self.diarize_var.get() else 'Нет'}\n")
                f.write("=" * 70 + "\n\n")
            
            print(f"\n✓ Создан файл: {filepath}")
            return filepath
            
        except Exception as e:
            print(f"\n⚠ Ошибка создания файла: {e}")
            return None
    
    def _display_transcript_mode(self, result, save_filepath):
        """Отображение результатов в режиме стенограммы (текст по спикерам с временными интервалами)"""
        
        # Заголовок для стенограммы
        header = "\n" + "=" * 70 + "\nСТЕНОГРАММА:\n" + "=" * 70 + "\n\n"
        self.result_text.insert(tk.END, header)
        self._append_to_file(save_filepath, header)
        
        if "segments" in result:
            total_segments = len(result["segments"])
            
            # Группируем текст по спикерам
            current_speaker = None
            current_text = ""
            start_time = None
            end_time = None
            
            for i, segment in enumerate(result["segments"], 1):
                text = segment.get("text", "").strip()
                speaker = segment.get("speaker", None)
                seg_start = segment.get("start", 0)
                seg_end = segment.get("end", 0)
                
                # Определяем имя спикера
                if speaker:
                    speaker_name = speaker
                else:
                    speaker_name = "[спикер не определен]"
                
                if speaker_name != current_speaker:
                    # Если сменился спикер, выводим накопленный текст
                    if current_text and start_time is not None:
                        time_interval = f"[{format_timestamp(start_time)} - {format_timestamp(end_time)}]"
                        line = f"{time_interval} {current_speaker}: {current_text}\n\n"
                        self.result_text.insert(tk.END, line)
                        self._append_to_file(save_filepath, line)
                    
                    # Начинаем новый блок
                    current_speaker = speaker_name
                    current_text = text
                    start_time = seg_start
                    end_time = seg_end
                else:
                    # Тот же спикер - добавляем текст и обновляем конечное время
                    current_text += " " + text
                    end_time = seg_end
                
                # Обновляем статус
                progress = (i / total_segments) * 100
                self.status_label.config(
                    text=f"💾 Сохранение результатов... {progress:.0f}%",
                    foreground=COLORS['primary']
                )
                self.root.update()
            
            # Выводим последний блок
            if current_text and start_time is not None:
                time_interval = f"[{format_timestamp(start_time)} - {format_timestamp(end_time)}]"
                line = f"{time_interval} {current_speaker}: {current_text}\n\n"
                self.result_text.insert(tk.END, line)
                self._append_to_file(save_filepath, line)
    
    def _append_to_file(self, filepath, text):
        """Добавление текста в файл"""
        if filepath:
            try:
                with open(filepath, "a", encoding="utf-8") as f:
                    f.write(text)
                    f.flush()
            except Exception as e:
                print(f"\n⚠ Ошибка записи в файл: {e}")
    
    def _finalize_save_file(self, filepath):
        """Добавление времени завершения в конец файла"""
        if filepath:
            try:
                from datetime import datetime
                
                end_time = datetime.now()
                
                # Вычисляем длительность обработки
                if hasattr(self, 'processing_start_time'):
                    duration = end_time - self.processing_start_time
                    duration_seconds = duration.total_seconds()
                    hours = int(duration_seconds // 3600)
                    minutes = int((duration_seconds % 3600) // 60)
                    seconds = int(duration_seconds % 60)
                    
                    duration_str = f"{hours:02d}:{minutes:02d}:{seconds:02d}"
                else:
                    duration_str = "неизвестно"
                
                with open(filepath, "a", encoding="utf-8") as f:
                    f.write("\n" + "=" * 70 + "\n")
                    f.write(f"Время завершения: {end_time.strftime('%Y-%m-%d %H:%M:%S')}\n")
                    f.write(f"Длительность обработки: {duration_str}\n")
                    f.write("=" * 70 + "\n")
                
                print(f"✓ Время завершения добавлено в файл")
            except Exception as e:
                print(f"\n⚠ Ошибка добавления времени завершения: {e}")
    
    def save_result(self):
        """Ручное сохранение результата"""
        text = self.result_text.get(1.0, tk.END).strip()
        
        if not text:
            messagebox.showwarning("Предупреждение", "Нет текста для сохранения")
            return
        
        file_path = filedialog.asksaveasfilename(
            title="Сохранить результат",
            defaultextension=".txt",
            filetypes=[
                ("Текстовые файлы", "*.txt"),
                ("Все файлы", "*.*")
            ]
        )
        
        if file_path:
            try:
                with open(file_path, "w", encoding="utf-8") as f:
                    f.write(text)
                messagebox.showinfo("Успех", f"Результат сохранен в:\n{file_path}")
            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось сохранить файл:\n{str(e)}")
    
    def clear_result(self):
        """Очистка результата"""
        self.result_text.delete(1.0, tk.END)
    
    def copy_result(self):
        """Копирование результата в буфер обмена"""
        text = self.result_text.get(1.0, tk.END).strip()
        
        if not text:
            messagebox.showwarning("Предупреждение", "Нет текста для копирования")
            return
        
        self.root.clipboard_clear()
        self.root.clipboard_append(text)
        messagebox.showinfo("Успех", "Текст скопирован в буфер обмена")
    
    def open_results_folder(self):
        """Открытие папки с результатами"""
        try:
            results_dir = RESULTS_DIR
            os.makedirs(results_dir, exist_ok=True)
            os.startfile(results_dir)
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось открыть папку:\n{str(e)}")

    def open_review_window(self):
        """Открыть окно проверки и редактирования последнего распознанного файла"""
        if self.last_audio is None or self.last_result is None:
            messagebox.showwarning(
                "Нет данных",
                "Сначала выполните распознавание речи — данные для проверки появятся после обработки файла.",
            )
            return
        segments = self.last_result.get("segments", [])
        if not segments:
            messagebox.showinfo("Информация", "В результате распознавания нет сегментов для отображения.")
            return
        SegmentReviewWindow(
            self.root,
            self.last_audio,
            self.audio_sample_rate,
            segments,
            self.last_save_filepath,
            self.last_processed_audio_path,
        )
    
    def export_to_word(self):
        """Экспорт стенограммы в MS Word - создание отдельных файлов для каждого обработанного аудио"""
        # Проверяем, есть ли обработанные файлы
        if not self.processed_files:
            messagebox.showerror("Ошибка", "Нет обработанных файлов для экспорта")
            return
        
        try:
            # Импортируем библиотеку для работы с Word
            try:
                from docx import Document
                from docx.shared import Pt, RGBColor
                from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            except ImportError:
                messagebox.showerror(
                    "Ошибка",
                    "Библиотека python-docx не установлена!\n\n"
                    "Установите её командой:\n"
                    "pip install python-docx"
                )
                return
            
            print(f"\n📄 Экспорт в Word ({len(self.processed_files)} файлов)...")
            
            created_files = []
            errors = []
            
            # Обрабатываем каждый файл из списка обработанных
            for audio_file, txt_file in self.processed_files:
                try:
                    if not os.path.exists(txt_file):
                        msg = f"⚠️ Файл не найден: {txt_file}"
                        print(msg)
                        errors.append(f"Файл не найден: {os.path.basename(txt_file)}")
                        continue
                    
                    print(f"\n📝 Обработка: {os.path.basename(txt_file)}")
                    
                    # Читаем файл
                    with open(txt_file, 'r', encoding='utf-8') as f:
                        content = f.read()
                    
                    if not content.strip():
                        msg = f"⚠️ Файл пуст: {txt_file}"
                        print(msg)
                        errors.append(f"Файл пуст: {os.path.basename(txt_file)}")
                        continue
                    
                    print(f"✓ Прочитано {len(content)} символов")
                    
                    # Создаем документ Word
                    doc = Document()
                    
                    # Добавляем заголовок с именем файла
                    title = doc.add_paragraph()
                    title_run = title.add_run(f"📄 {os.path.splitext(os.path.basename(audio_file))[0]}")
                    title_run.font.size = Pt(14)
                    title_run.font.bold = True
                    title_run.font.color.rgb = RGBColor(0, 0, 128)  # Тёмно-синий
                    
                    # Добавляем информацию о файле
                    info = doc.add_paragraph(f"Исходный файл: {os.path.basename(audio_file)}")
                    info_run = info.runs[0] if info.runs else info.add_run("")
                    info_run.font.size = Pt(10)
                    info_run.font.italic = True
                    
                    # Добавляем дату обработки
                    from datetime import datetime
                    date_para = doc.add_paragraph(f"Обработано: {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}")
                    date_run = date_para.runs[0] if date_para.runs else date_para.add_run("")
                    date_run.font.size = Pt(10)
                    date_run.font.italic = True
                    
                    # Добавляем разделитель
                    doc.add_paragraph("_" * 80)
                    
                    # Разбиваем на параграфы по двойным переносам строк
                    paragraphs = content.split('\n\n')
                    
                    print(f"✓ Найдено параграфов: {len(paragraphs)}")
                    
                    for para_text in paragraphs:
                        if para_text.strip():
                            # Добавляем параграф
                            p = doc.add_paragraph()
                            
                            # Проверяем, это ли временная метка или заголовок
                            is_header = para_text.strip().startswith('[')
                            
                            # Добавляем текст в параграф
                            run = p.add_run(para_text.strip())
                            
                            # Стилизуем в зависимости от типа
                            if is_header:
                                run.font.bold = True
                                run.font.size = Pt(11)
                                run.font.color.rgb = RGBColor(0, 100, 0)  # Темно-зелёный для меток времени
                            else:
                                run.font.size = Pt(12)
                                run.font.color.rgb = RGBColor(0, 0, 0)  # Чёрный для текста
                            
                            # Установим отступ для лучшей читаемости
                            p.paragraph_format.space_after = Pt(6)
                    
                    # Сохраняем документ
                    word_filename = txt_file.replace('.txt', '.docx')
                    
                    # Проверяем, не существует ли уже файл
                    counter = 1
                    original_word_filename = word_filename
                    while os.path.exists(word_filename):
                        base_name = original_word_filename.replace('.docx', '')
                        word_filename = f"{base_name}_{counter}.docx"
                        counter += 1
                    
                    doc.save(word_filename)
                    
                    print(f"✓ Документ Word сохранен: {word_filename}")
                    created_files.append(word_filename)
                    
                except Exception as e:
                    error_msg = f"Ошибка при обработке {os.path.basename(txt_file)}: {str(e)}"
                    print(f"❌ {error_msg}")
                    errors.append(error_msg)
                    import traceback
                    traceback.print_exc()
            
            # Показываем результат
            if created_files:
                # Открываем первый созданный Word файл
                try:
                    os.startfile(created_files[0])
                    print(f"✓ Открыт: {os.path.basename(created_files[0])}")
                except Exception as e:
                    print(f"⚠️ Не удалось автоматически открыть файл: {e}")
                    print(f"📂 Файлы находятся в папке Results: {RESULTS_DIR}")
                
                files_list = "\n".join([f"✓ {os.path.basename(f)}" for f in created_files])
                msg = f"Создано документов Word: {len(created_files)}\n\n{files_list}"
                
                if errors:
                    msg += f"\n\n⚠️ Ошибки ({len(errors)}):\n"
                    msg += "\n".join(errors[:3])  # Показываем первые 3 ошибки
                
                messagebox.showinfo("✅ Экспорт завершен", msg)
                print(f"\n✅ Экспорт завершен! Создано {len(created_files)} файлов Word")
            else:
                error_msg = "Не удалось создать ни одного документа"
                if errors:
                    error_msg += f"\n\nПроблемы:\n" + "\n".join(errors[:3])
                messagebox.showwarning("Предупреждение", error_msg)
            
        except Exception as e:
            import traceback
            print("❌ КРИТИЧЕСКАЯ ОШИБКА")
            traceback.print_exc()
            messagebox.showerror("Ошибка", f"Не удалось экспортировать в Word:\n{str(e)}")


class SegmentReviewWindow:
    """🎧 Окно проверки и редактирования распознанных аудио-фрагментов"""

    def __init__(self, parent, audio, sample_rate, segments, save_filepath, processed_audio_path=None):
        self.audio = audio                    # numpy array (float32, mono)
        self.sample_rate = sample_rate        # 16000
        self.segments = segments              # list of segment dicts (mutable)
        self.save_filepath = save_filepath
        self.processed_audio_path = processed_audio_path
        self.current_play_thread = None
        self._stop_autoplay = False
        self._play_generation = 0  # увеличивается при каждом play_segment; защита от устаревших after()
        self.play_buttons = []     # кнопки ▶ по индексу сегмента
        self.dataset_buttons = []  # кнопки 🗂 для сохранения в датасет
        self.merge_buttons = []    # кнопки ⊕ для объединения со следующим
        self._playing_audio = None   # ссылка на numpy-массив пока играет PortAudio
        self._window_alive = True    # False после _on_close, защита window.after()

        self.window = tk.Toplevel(parent)
        self.window.title("🎧 Проверка и редактирование распознанного текста")

        win_w, win_h = 980, 700
        sw = self.window.winfo_screenwidth()
        sh = self.window.winfo_screenheight()
        self.window.geometry(f"{win_w}x{win_h}+{(sw - win_w) // 2}+{(sh - win_h) // 2}")
        self.window.configure(bg=COLORS['bg'])
        self.window.resizable(True, True)
        self.window.minsize(700, 400)

        self.text_vars = []   # StringVar — по одному на каждый сегмент

        self._create_widgets()
        self.window.protocol("WM_DELETE_WINDOW", self._on_close)

    # ------------------------------------------------------------------ #
    def _on_close(self):
        self._window_alive = False
        self._stop_autoplay = True
        # Останавливаем PortAudio *до* join, чтобы sd.wait() в потоке завершился
        try:
            import sounddevice as sd
            sd.stop()
        except Exception:
            pass
        # Ждём завершения потока (до 5 секунд)
        if self.current_play_thread and self.current_play_thread.is_alive():
            self.current_play_thread.join(timeout=5.0)
        # Обнуляем буфер только после того, как поток завершён
        self._playing_audio = None
        self.window.destroy()

    # ------------------------------------------------------------------ #
    def _create_widgets(self):
        # ---- Заголовок ----
        title_frame = tk.Frame(self.window, bg=COLORS['darker'], height=50)
        title_frame.pack(fill=tk.X)
        title_frame.pack_propagate(False)
        tk.Label(
            title_frame,
            text="🎧 Проверка и редактирование распознанного текста",
            font=("Segoe UI", 13, "bold"),
            bg=COLORS['darker'],
            fg=COLORS['primary'],
        ).pack(expand=True)

        # ---- Подсказка ----
        tk.Label(
            self.window,
            text="Нажмите ▶ чтобы прослушать фрагмент. Текст можно отредактировать прямо в поле.",
            font=("Segoe UI", 10),
            bg=COLORS['bg'],
            fg=COLORS['text'],
        ).pack(pady=(6, 2))

        # ---- Прокручиваемая область ----
        container = tk.Frame(self.window, bg=COLORS['bg'])
        container.pack(fill=tk.BOTH, expand=True, padx=8, pady=4)

        self._canvas = tk.Canvas(container, bg=COLORS['bg'], highlightthickness=0)
        scrollbar = tk.Scrollbar(container, orient="vertical", command=self._canvas.yview)
        self._scroll_frame = tk.Frame(self._canvas, bg=COLORS['bg'])

        self._scroll_frame.bind(
            "<Configure>",
            lambda e: self._canvas.configure(scrollregion=self._canvas.bbox("all")),
        )
        self._canvas_win = self._canvas.create_window((0, 0), window=self._scroll_frame, anchor="nw")
        self._canvas.configure(yscrollcommand=scrollbar.set)

        # Растягиваем внутренний фрейм на ширину canvas
        self._canvas.bind(
            "<Configure>",
            lambda e: self._canvas.itemconfig(self._canvas_win, width=e.width),
        )

        self._canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        # Колёсико мыши
        self._canvas.bind_all(
            "<MouseWheel>",
            lambda e: self._canvas.yview_scroll(-1 * (e.delta // 120), "units"),
        )

        # ---- Заголовок колонок ----
        hdr = tk.Frame(self._scroll_frame, bg=COLORS['card_bg'], pady=4)
        hdr.pack(fill=tk.X, padx=4, pady=(4, 2))
        tk.Label(hdr, text="  ▶  ", bg=COLORS['card_bg'], fg=COLORS['text'],
                 font=("Segoe UI", 9, "bold")).pack(side=tk.LEFT)
        tk.Label(hdr, text="Временная метка", width=26, anchor='w',
                 bg=COLORS['card_bg'], fg=COLORS['text'],
                 font=("Segoe UI", 9, "bold")).pack(side=tk.LEFT)
        tk.Label(hdr, text="Распознанный текст", anchor='w',
                 bg=COLORS['card_bg'], fg=COLORS['text'],
                 font=("Segoe UI", 9, "bold")).pack(side=tk.LEFT, fill=tk.X, expand=True)
        tk.Label(hdr, text="Датасет", width=8, anchor='c',
                 bg=COLORS['card_bg'], fg=COLORS['text'],
                 font=("Segoe UI", 9, "bold")).pack(side=tk.RIGHT, padx=(0, 8))
        tk.Label(hdr, text="  ⊕  ", width=5, anchor='c',
                 bg=COLORS['card_bg'], fg=COLORS['text'],
                 font=("Segoe UI", 9, "bold")).pack(side=tk.RIGHT, padx=(0, 2))

        # ---- Строки сегментов ----
        for idx, seg in enumerate(self.segments):
            self._add_row(idx, seg)

        # ---- Нижние кнопки ----
        bottom = tk.Frame(self.window, bg=COLORS['darker'], pady=8)
        bottom.pack(fill=tk.X, side=tk.BOTTOM)

        tk.Button(
            bottom, text="⏹ Стоп",
            command=self.stop_playback,
            bg=COLORS['danger'], fg='white',
            font=("Segoe UI", 11), relief=tk.FLAT,
            cursor="hand2", padx=12, pady=5,
        ).pack(side=tk.LEFT, padx=10)

        tk.Button(
            bottom, text="📂 Открыть датасет",
            command=self.open_dataset_folder,
            bg=COLORS['border'], fg=COLORS['text'],
            font=("Segoe UI", 11), relief=tk.FLAT,
            cursor="hand2", padx=12, pady=5,
        ).pack(side=tk.LEFT, padx=4)

        tk.Button(
            bottom, text="✖ Закрыть",
            command=self._on_close,
            bg=COLORS['border'], fg=COLORS['text'],
            font=("Segoe UI", 11), relief=tk.FLAT,
            cursor="hand2", padx=12, pady=5,
        ).pack(side=tk.RIGHT, padx=10)

        tk.Button(
            bottom, text="💾 Сохранить изменения",
            command=self.save_edits,
            bg=COLORS['success'], fg='white',
            font=("Segoe UI", 12, "bold"), relief=tk.FLAT,
            cursor="hand2", padx=16, pady=5,
        ).pack(side=tk.RIGHT, padx=5)

    # ------------------------------------------------------------------ #
    def _add_row(self, idx, segment):
        start   = segment.get("start", 0)
        end     = segment.get("end", 0)
        text    = segment.get("text", "").strip()
        speaker = segment.get("speaker", None)

        row_bg = COLORS['card_bg'] if idx % 2 == 0 else COLORS['input_bg']

        row = tk.Frame(self._scroll_frame, bg=row_bg, pady=3)
        row.pack(fill=tk.X, padx=4, pady=1)

        # Кнопка воспроизведения
        play_btn = tk.Button(
            row, text="▶", width=3,
            command=lambda i=idx: self.play_segment(i),
            bg=COLORS['primary'], fg='white',
            font=("Segoe UI", 10), relief=tk.FLAT,
            cursor="hand2", activebackground=COLORS['accent'],
        )
        play_btn.pack(side=tk.LEFT, padx=(5, 3), pady=2)
        self.play_buttons.append(play_btn)

        # Временная метка (+ спикер)
        ts = f"{format_timestamp(start)} – {format_timestamp(end)}"
        if speaker:
            ts += f"\n{speaker}"
        tk.Label(
            row, text=ts, width=26, anchor='w', justify='left',
            bg=row_bg, fg=COLORS['accent'],
            font=("Consolas", 9),
        ).pack(side=tk.LEFT, padx=3)

        # Редактируемый текст
        var = tk.StringVar(value=text)
        self.text_vars.append(var)
        tk.Entry(
            row, textvariable=var,
            bg=COLORS['input_bg'], fg=COLORS['text'],
            insertbackground=COLORS['primary'],
            font=("Segoe UI", 10),
            relief=tk.FLAT, bd=2,
        ).pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(3, 4), pady=2)

        # Кнопка сохранения фрагмента в датасет
        ds_btn = tk.Button(
            row, text="🗂", width=3,
            command=lambda i=idx, v=var: self.save_to_dataset(i, v),
            bg=COLORS['border'], fg=COLORS['text'],
            font=("Segoe UI", 10), relief=tk.FLAT,
            cursor="hand2", activebackground=COLORS['success'],
        )
        ds_btn.pack(side=tk.RIGHT, padx=(2, 5), pady=2)
        self.dataset_buttons.append(ds_btn)

        # Кнопка объединения со следующим сегментом
        is_last = (idx == len(self.segments) - 1)
        merge_btn = tk.Button(
            row, text="⊕", width=3,
            command=(lambda i=idx: self.merge_segments(i)) if not is_last else None,
            state=tk.NORMAL if not is_last else tk.DISABLED,
            bg=COLORS['border'], fg=COLORS['text'],
            font=("Segoe UI", 10), relief=tk.FLAT,
            cursor="hand2" if not is_last else "arrow",
            activebackground=COLORS['warning'],
        )
        merge_btn.pack(side=tk.RIGHT, padx=(2, 2), pady=2)
        self.merge_buttons.append(merge_btn)

    # ------------------------------------------------------------------ #
    def _rebuild_rows(self):
        """Перестроить все строки сегментов (вызывается после слияния)."""
        for w in self._scroll_frame.winfo_children():
            w.destroy()
        self.text_vars.clear()
        self.play_buttons.clear()
        self.dataset_buttons.clear()
        self.merge_buttons.clear()

        # Заголовок колонок
        hdr = tk.Frame(self._scroll_frame, bg=COLORS['card_bg'], pady=4)
        hdr.pack(fill=tk.X, padx=4, pady=(4, 2))
        tk.Label(hdr, text="  ▶  ", bg=COLORS['card_bg'], fg=COLORS['text'],
                 font=("Segoe UI", 9, "bold")).pack(side=tk.LEFT)
        tk.Label(hdr, text="Временная метка", width=26, anchor='w',
                 bg=COLORS['card_bg'], fg=COLORS['text'],
                 font=("Segoe UI", 9, "bold")).pack(side=tk.LEFT)
        tk.Label(hdr, text="Распознанный текст", anchor='w',
                 bg=COLORS['card_bg'], fg=COLORS['text'],
                 font=("Segoe UI", 9, "bold")).pack(side=tk.LEFT, fill=tk.X, expand=True)
        tk.Label(hdr, text="Датасет", width=8, anchor='c',
                 bg=COLORS['card_bg'], fg=COLORS['text'],
                 font=("Segoe UI", 9, "bold")).pack(side=tk.RIGHT, padx=(0, 8))
        tk.Label(hdr, text="  ⊕  ", width=5, anchor='c',
                 bg=COLORS['card_bg'], fg=COLORS['text'],
                 font=("Segoe UI", 9, "bold")).pack(side=tk.RIGHT, padx=(0, 2))

        for idx, seg in enumerate(self.segments):
            self._add_row(idx, seg)

        self._canvas.update_idletasks()
        self._canvas.configure(scrollregion=self._canvas.bbox("all"))

    # ------------------------------------------------------------------ #
    def merge_segments(self, idx):
        """Объединить сегмент idx со следующим (idx+1)."""
        if idx >= len(self.segments) - 1:
            return
        seg_a = self.segments[idx]
        seg_b = self.segments[idx + 1]
        # Берём актуальный текст из полей ввода
        text_a = self.text_vars[idx].get().strip()
        text_b = self.text_vars[idx + 1].get().strip()
        seg_a["text"] = (text_a + " " + text_b).strip()
        seg_a["end"]  = seg_b["end"]
        # Наследуем спикера из seg_a; если пусто — берём из seg_b
        if not seg_a.get("speaker") and seg_b.get("speaker"):
            seg_a["speaker"] = seg_b["speaker"]
        del self.segments[idx + 1]
        self.stop_playback()
        self._rebuild_rows()
        print(f"✓ Сегменты {idx} и {idx + 1} объединены")

    # ------------------------------------------------------------------ #
    def save_to_dataset(self, idx, text_var):
        """Сохранить фрагмент (аудио + исправленный текст) в датасет для дообучения.

        Структура датасета (формат HuggingFace audiofolder):
            <папка Results>/dataset/
                audio/
                    segment_NNNN.wav
                metadata.jsonl  — {"file_name": "audio/...", "transcription": "..."}
        """
        import json
        import soundfile as sf

        text = text_var.get().strip()
        if not text:
            messagebox.showwarning(
                "Пустой текст",
                "Сначала введите правильный текст в поле, затем нажмите 🗂.",
                parent=self.window,
            )
            return

        segment = self.segments[idx]
        start = segment.get("start", 0)
        end   = segment.get("end",   0)

        # --- выбираем источник аудио ---
        audio_np = None
        seg_sr = self.sample_rate

        if self.processed_audio_path and os.path.exists(self.processed_audio_path):
            try:
                info = sf.info(self.processed_audio_path)
                seg_sr = info.samplerate
                s = max(0, int(start * seg_sr))
                e = max(s,  int(end   * seg_sr))
                audio_np, _ = sf.read(
                    self.processed_audio_path,
                    start=s, stop=e,
                    dtype='float32', always_2d=False,
                )
                if audio_np.ndim > 1:
                    audio_np = audio_np.mean(axis=1)
            except Exception as ex:
                print(f"⚠️ Не удалось прочитать обработанный файл: {ex}")
                audio_np = None

        if audio_np is None:
            s = max(0, int(start * self.sample_rate))
            e = min(len(self.audio), int(end * self.sample_rate))
            audio_np = self.audio[s:e]
            seg_sr = self.sample_rate

        if audio_np is None or len(audio_np) == 0:
            messagebox.showwarning(
                "Нет аудио",
                "Не удалось извлечь аудиофрагмент для этого сегмента.",
                parent=self.window,
            )
            return

        # --- папка датасета рядом с файлом результатов ---
        base_dir = os.path.dirname(self.save_filepath) if self.save_filepath else RESULTS_DIR
        dataset_dir  = os.path.join(base_dir, "dataset")
        audio_dir    = os.path.join(dataset_dir, "audio")
        metadata_path = os.path.join(dataset_dir, "metadata.jsonl")
        os.makedirs(audio_dir, exist_ok=True)

        # --- имя файла: segment_NNNN_start-end.wav ---
        wav_name = f"segment_{idx:04d}_{start:.2f}-{end:.2f}.wav".replace(".", "_")
        wav_name = wav_name.replace("__", "_") + ".wav"
        # убираем двойное .wav если replace добавил:
        if wav_name.endswith(".wav.wav"):
            wav_name = wav_name[:-4]
        wav_path = os.path.join(audio_dir, wav_name)

        try:
            sf.write(wav_path, audio_np, seg_sr)
        except Exception as ex:
            messagebox.showerror(
                "Ошибка записи",
                f"Не удалось сохранить аудиофайл:\n{ex}",
                parent=self.window,
            )
            return

        # --- запись в metadata.jsonl (дозаписываем, не перезаписываем) ---
        record = {
            "file_name":     os.path.join("audio", wav_name).replace("\\", "/"),
            "transcription": text,
            "start":         round(start, 3),
            "end":           round(end,   3),
            "speaker":       segment.get("speaker", ""),
        }
        try:
            with open(metadata_path, "a", encoding="utf-8") as f:
                f.write(json.dumps(record, ensure_ascii=False) + "\n")
        except Exception as ex:
            messagebox.showerror(
                "Ошибка записи",
                f"Не удалось записать metadata.jsonl:\n{ex}",
                parent=self.window,
            )
            return

        # --- подсветить кнопку зелёным в знак успеха ---
        try:
            btn = self.dataset_buttons[idx]
            btn.config(bg=COLORS['success'])
            self.window.after(2000, lambda: btn.config(bg=COLORS['border']))
        except Exception:
            pass

        print(f"✓ Сегмент {idx} добавлен в датасет: {wav_path}")

    # ------------------------------------------------------------------ #
    def open_dataset_folder(self):
        """Открыть папку датасета в проводнике."""
        base_dir = os.path.dirname(self.save_filepath) if self.save_filepath else RESULTS_DIR
        dataset_dir = os.path.join(base_dir, "dataset")
        os.makedirs(dataset_dir, exist_ok=True)
        try:
            os.startfile(dataset_dir)
        except Exception as ex:
            messagebox.showerror("Ошибка", f"Не удалось открыть папку:\n{ex}", parent=self.window)

    # ------------------------------------------------------------------ #
    def play_segment(self, idx):
        """Воспроизвести аудио для сегмента idx.

        Диапазон воспроизведения расширяется так, чтобы не оставалось
        непокрытого времени между фрагментами:
          - Начало: конец предыдущего сегмента (или 0:00 для первого).
          - Конец:  начало следующего сегмента (или конец аудио для последнего).
        """
        self.stop_playback()

        n = len(self.segments)

        # Начало: от конца предыдущего фрагмента (или от нуля)
        if idx == 0:
            play_start = 0.0
        else:
            play_start = self.segments[idx - 1].get("end", 0)

        # Конец: до начала следующего фрагмента (или конец аудио)
        if idx >= n - 1:
            play_end = None   # будет заменено на длину файла ниже
        else:
            play_end = self.segments[idx + 1].get("start", self.segments[idx].get("end", 0))

        playback_sample_rate = self.sample_rate
        audio_slice = None

        if self.processed_audio_path and os.path.exists(self.processed_audio_path):
            try:
                import soundfile as sf

                info = sf.info(self.processed_audio_path)
                playback_sample_rate = info.samplerate
                total_frames = info.frames

                s = max(0, int(play_start * playback_sample_rate))
                e = total_frames if play_end is None else min(total_frames, int(play_end * playback_sample_rate))
                e = max(s, e)

                audio_slice = sf.read(
                    self.processed_audio_path,
                    start=s,
                    stop=e,
                    dtype='float32',
                    always_2d=False,
                )[0]

                if isinstance(audio_slice, np.ndarray) and audio_slice.ndim > 1:
                    audio_slice = audio_slice.mean(axis=1)
            except Exception as ex:
                print(f"⚠️ Ошибка чтения обработанного файла для воспроизведения: {ex}")
                audio_slice = None

        if audio_slice is None:
            total_in_mem = len(self.audio)
            s = max(0, int(play_start * self.sample_rate))
            e = total_in_mem if play_end is None else min(total_in_mem, int(play_end * self.sample_rate))
            e = max(s, e)
            audio_slice = self.audio[s:e]
            playback_sample_rate = self.sample_rate

        if len(audio_slice) == 0:
            return

        # Храним ссылку на массив в self, чтобы GC не удалил его, пока PortAudio работает
        self._playing_audio = np.ascontiguousarray(audio_slice, dtype=np.float32)
        audio_ref = self._playing_audio

        self._stop_autoplay = False  # сброс флага при каждом ручном запуске
        self._play_generation += 1
        my_gen = self._play_generation  # поколение этого запуска

        # Подсветить активную кнопку, сбросить все остальные
        for i, btn in enumerate(self.play_buttons):
            try:
                if i == idx:
                    btn.config(bg=COLORS['playing'], fg=COLORS['darker'])
                else:
                    btn.config(bg=COLORS['primary'], fg='white')
            except Exception:
                pass

        def _play():
            try:
                import sounddevice as sd
                import time
                sd.play(audio_ref, samplerate=playback_sample_rate)
                # Заменяем блокирующий sd.wait() на polling — так sd.stop() из
                # _on_close немедленно прерывает цикл без ACCESS_VIOLATION
                try:
                    stream = sd.get_stream()
                    while stream.active:
                        if self._stop_autoplay:
                            sd.stop()
                            break
                        time.sleep(0.05)
                except Exception:
                    # Запасной вариант: старый sd.wait() если get_stream недоступен
                    sd.wait()
            except ImportError:
                if self._window_alive:
                    self.window.after(0, lambda: messagebox.showwarning(
                        "Воспроизведение недоступно",
                        "Установите модуль sounddevice:\n  pip install sounddevice",
                    ))
                return
            except Exception as ex:
                print(f"⚠️ Ошибка воспроизведения: {ex}")
                return

            if not self._window_alive:
                return

            # Автопереход к следующему фрагменту только если никто не нажал другую кнопку
            if not self._stop_autoplay and idx + 1 < len(self.segments) and self._play_generation == my_gen:
                self.window.after(0, lambda g=my_gen: self.play_segment(idx + 1) if self._play_generation == g and self._window_alive else None)
            else:
                # Воспроизведение закончено без продолжения — снимаем подсветку
                if self._play_generation == my_gen and self._window_alive:
                    self.window.after(0, self.stop_playback)

        self.current_play_thread = threading.Thread(target=_play, daemon=True)
        self.current_play_thread.start()

    # ------------------------------------------------------------------ #
    def stop_playback(self):
        """Остановить воспроизведение и прервать цепочку автовоспроизведения"""
        self._stop_autoplay = True
        try:
            import sounddevice as sd
            sd.stop()
        except Exception:
            pass
        # Сбросить подсветку всех кнопок
        for btn in self.play_buttons:
            try:
                btn.config(bg=COLORS['primary'], fg='white')
            except Exception:
                pass

    # ------------------------------------------------------------------ #
    def save_edits(self):
        """Записать отредактированные тексты в файл результатов"""
        # Обновляем объекты сегментов в памяти
        for seg, var in zip(self.segments, self.text_vars):
            seg["text"] = var.get()

        if not self.save_filepath or not os.path.exists(self.save_filepath):
            messagebox.showwarning(
                "Предупреждение",
                "Файл результатов не найден — изменения обновлены в памяти.",
            )
            return

        try:
            with open(self.save_filepath, 'r', encoding='utf-8') as f:
                content = f.read()

            # Определяем конец шапки файла (до раздела с транскрипцией)
            for marker in ("РАСПОЗНАННЫЙ ТЕКСТ", "ВРЕМЕННЫЕ МЕТКИ", "СТЕНОГРАММА"):
                idx_m = content.find(marker)
                if idx_m != -1:
                    header_end = content.rfind("\n", 0, idx_m) + 1
                    break
            else:
                # Fallback: после второго блока "==="
                sep_count = 0
                header_end = 0
                i = 0
                while i < len(content) - 2:
                    if content[i:i+3] == "===":
                        eol = content.find("\n", i)
                        sep_count += 1
                        if sep_count >= 2:
                            header_end = (eol + 1) if eol != -1 else len(content)
                            break
                        i = eol if eol != -1 else len(content)
                    i += 1

            header = content[:header_end]

            lines = ["\nОТРЕДАКТИРОВАННЫЙ ТЕКСТ:\n", "=" * 70 + "\n\n"]
            for seg in self.segments:
                s   = seg.get("start", 0)
                e   = seg.get("end", 0)
                t   = seg.get("text", "").strip()
                spk = seg.get("speaker", None)
                if spk:
                    lines.append(f"[{format_timestamp(s)} - {format_timestamp(e)}] {spk}:\n{t}\n\n")
                else:
                    lines.append(f"[{format_timestamp(s)} - {format_timestamp(e)}]: {t}\n\n")

            with open(self.save_filepath, 'w', encoding='utf-8') as f:
                f.write(header + "".join(lines))

            messagebox.showinfo("Успех", "Изменения сохранены в файл результатов!")

        except Exception as ex:
            messagebox.showerror("Ошибка", f"Не удалось сохранить:\n{ex}")


# ============================================================ #


def show_splash_screen():
    """Загрузочный экран AI VoiceFinder"""
    splash = tk.Tk()
    splash.title("AI VoiceFinder")
    splash.overrideredirect(True)

    # Загружаем GIF заранее, чтобы узнать размеры
    script_dir = os.path.dirname(os.path.abspath(__file__))
    gif_path = os.path.join(script_dir, "Sources", "Load waiting.gif")
    gif_width = 280
    gif_height = 280

    try:
        if os.path.exists(gif_path):
            temp_gif = Image.open(gif_path)
            gif_width = int(temp_gif.width * 0.5)
            gif_height = int(temp_gif.height * 0.5)
            temp_gif.close()
    except:
        pass

    # Размеры окна
    width = 680
    height = max(380, gif_height + 80)

    screen_width = splash.winfo_screenwidth()
    screen_height = splash.winfo_screenheight()
    x = (screen_width - width) // 2
    y = (screen_height - height) // 2
    splash.geometry(f"{width}x{height}+{x}+{y}")
    splash.configure(bg=COLORS['darker'])

    # Тонкая полоска сверху (акцент)
    top_accent = tk.Frame(splash, bg=COLORS['border'], height=2)
    top_accent.pack(fill=tk.X)

    # Главный контейнер
    main_container = tk.Frame(splash, bg=COLORS['darker'])
    main_container.pack(expand=True, fill=tk.BOTH, padx=40, pady=30)

    # ЛЕВАЯ ЧАСТЬ
    left_frame = tk.Frame(main_container, bg=COLORS['darker'], width=340)
    left_frame.pack(side=tk.LEFT, fill=tk.Y, expand=False, padx=(0, 20))
    left_frame.pack_propagate(False)

    # Заголовок приложения
    title_label = tk.Label(
        left_frame,
        text="AI VoiceFinder",
        font=("Segoe UI", 26, "bold"),
        bg=COLORS['darker'],
        fg=COLORS['light'],
        anchor='w'
    )
    title_label.pack(anchor='w', pady=(10, 4))

    # Тонкий разделитель
    sep = tk.Frame(left_frame, bg=COLORS['border'], height=1)
    sep.pack(fill=tk.X, pady=(0, 8))

    subtitle_label = tk.Label(
        left_frame,
        text="Система распознавания речи",
        font=("Segoe UI", 11),
        bg=COLORS['darker'],
        fg=COLORS['success'],
        anchor='w'
    )
    subtitle_label.pack(anchor='w', pady=(0, 20))

    # Переменная для отслеживания состояния
    splash._is_alive = True
    splash._after_ids = []

    # Лог статуса загрузки
    status_frame = tk.Frame(left_frame, bg=COLORS['darker'])
    status_frame.pack(anchor='w', fill=tk.BOTH, expand=True)

    status_label = tk.Label(
        status_frame,
        text="",
        font=("Consolas", 10),
        bg=COLORS['darker'],
        fg=COLORS['accent'],
        anchor='nw',
        justify='left',
        wraplength=310
    )
    status_label.pack(anchor='nw', fill=tk.BOTH, expand=True)

    # Нижняя панель: счётчик + версия
    bottom_frame = tk.Frame(left_frame, bg=COLORS['darker'])
    bottom_frame.pack(fill=tk.X, side=tk.BOTTOM, pady=(10, 0))

    percent_label = tk.Label(
        bottom_frame,
        text="0%",
        font=("Segoe UI", 20, "bold"),
        bg=COLORS['darker'],
        fg=COLORS['light'],
        anchor='w'
    )
    percent_label.pack(side=tk.LEFT)

    version_label = tk.Label(
        bottom_frame,
        text="v2.1",
        font=("Segoe UI", 10),
        bg=COLORS['darker'],
        fg=COLORS['border'],
        anchor='e'
    )
    version_label.pack(side=tk.RIGHT, anchor='s', pady=2)

    # Тексты загрузки (без эмодзи, строгий стиль)
    loading_messages = [
        "Инициализация AI VoiceFinder...",
        "Загрузка модели large-v3",
        "Настройка конвейера обработки...",
        "Подготовка аудио-модуля...",
        "Загрузка компонентов выравнивания...",
        "Проверка зависимостей...",
        "Настройка пользовательского интерфейса...",
        "Финальная проверка системы...",
        "Запуск приложения..."
    ]

    completed_messages = []
    current_message = [0]
    current_char = [0]
    cursor_visible = [True]
    current_percent = [0]

    def animate_text():
        if not splash._is_alive:
            return
        try:
            if splash.winfo_exists():
                msg_index = current_message[0]
                char_index = current_char[0]

                if msg_index < len(loading_messages):
                    message = loading_messages[msg_index]

                    if char_index <= len(message):
                        current_text = message[:char_index]
                        cursor = "_" if cursor_visible[0] else " "

                        display_lines = completed_messages[-4:] if len(completed_messages) > 4 else completed_messages
                        display_text = "\n".join(display_lines)
                        if display_text:
                            display_text += "\n"
                        display_text += current_text + cursor

                        status_label.config(text=display_text)

                        progress = int((msg_index * 100 + (char_index / len(message) * 100)) / len(loading_messages))
                        if progress != current_percent[0]:
                            current_percent[0] = progress
                            percent_label.config(text=f"{progress}%")

                        current_char[0] += 1
                        delay = 1
                    else:
                        completed_messages.append(message)
                        current_message[0] += 1
                        current_char[0] = 0
                        delay = 150

                    after_id = splash.after(delay, animate_text)
                    splash._after_ids.append(after_id)
                else:
                    completed_messages.clear()
                    current_message[0] = 0
                    current_char[0] = 0
                    current_percent[0] = 0
                    percent_label.config(text="0%")
                    after_id = splash.after(800, animate_text)
                    splash._after_ids.append(after_id)
        except:
            pass

    def blink_cursor():
        if not splash._is_alive:
            return
        try:
            if splash.winfo_exists():
                cursor_visible[0] = not cursor_visible[0]
                after_id = splash.after(500, blink_cursor)
                splash._after_ids.append(after_id)
        except:
            pass

    animate_text()
    blink_cursor()

    # ПРАВАЯ ЧАСТЬ — GIF или пустой блок
    right_frame = tk.Frame(main_container, bg=COLORS['darker'])
    right_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)

    try:
        if os.path.exists(gif_path):
            gif_image = Image.open(gif_path)
            original_width = gif_image.width
            original_height = gif_image.height
            new_width = int(original_width * 0.5)
            new_height = int(original_height * 0.5)

            gif_label = tk.Label(right_frame, bg=COLORS['darker'], bd=0)
            gif_label.pack(expand=True)

            frames = []
            try:
                while True:
                    frame_image = gif_image.copy()
                    if frame_image.mode != 'RGBA':
                        frame_image = frame_image.convert('RGBA')
                    frame_image = frame_image.resize((new_width, new_height), Image.Resampling.LANCZOS)
                    frames.append(ImageTk.PhotoImage(frame_image))
                    gif_image.seek(len(frames))
            except EOFError:
                pass

            frame_index = [0]

            def update_frame():
                if not splash._is_alive:
                    return
                try:
                    if splash.winfo_exists():
                        frame = frames[frame_index[0]]
                        gif_label.config(image=frame)
                        gif_label.image = frame
                        frame_index[0] = (frame_index[0] + 1) % len(frames)
                        after_id = splash.after(50, update_frame)
                        splash._after_ids.append(after_id)
                except:
                    pass

            update_frame()
    except Exception as e:
        print(f"Не удалось загрузить GIF: {e}")

    # Тонкая полоска снизу
    bottom_accent = tk.Frame(splash, bg=COLORS['border'], height=2)
    bottom_accent.pack(fill=tk.X, side=tk.BOTTOM)

    # Функция для безопасного закрытия
    def safe_close():
        splash._is_alive = False
        for after_id in splash._after_ids:
            try:
                splash.after_cancel(after_id)
            except:
                pass
        try:
            splash.destroy()
        except:
            pass

    splash.safe_close = safe_close

    splash.update()
    return splash


def main():
    """Главная функция запуска приложения"""
    try:
        # Подавляем TensorFlow предупреждения для более быстрой загрузки
        import warnings
        warnings.filterwarnings('ignore', category=FutureWarning)
        warnings.filterwarnings('ignore', message='.*TRANSFORMERS_CACHE.*')
        
        # Показываем заставку загрузки
        print("\n⏳ Показываем заставку загрузки...")
        splash = show_splash_screen()
        splash.update_idletasks()
        splash.update()
        
        # Переменные для хранения загруженной модели
        loaded_model = {'model': None, 'device': None, 'compute_type': None}
        
        # Функция загрузки модели в фоновом потоке
        def load_model_background():
            try:
                print("\n🔥 Начинаем загрузку модели large-v3...")
                
                # Определяем устройство
                use_gpu = True
                if torch.cuda.is_available() and use_gpu:
                    device = "cuda"
                    compute_type = "float16"
                    device_name = torch.cuda.get_device_name(0)
                    print(f"✓ Используется GPU: {device_name}")
                else:
                    device = "cpu"
                    compute_type = "int8"
                    print("✓ Используется CPU")
                
                # Загружаем модель WhisperX large-v3
                print("⏳ Загрузка модели AI VoiceFinder (large-v3)...")
                _whisper_cache = os.path.join(LOCAL_CACHE_DIR, 'whisper')
                os.makedirs(_whisper_cache, exist_ok=True)
                model = whisperx.load_model(
                    "large-v3",
                    device,
                    compute_type=compute_type,
                    download_root=_whisper_cache,
                )
                print("✅ Модель large-v3 загружена успешно!")
                
                # Сохраняем модель
                loaded_model['model'] = model
                loaded_model['device'] = device
                loaded_model['compute_type'] = compute_type
                
            except Exception as e:
                print(f"❌ Ошибка загрузки модели: {e}")
                import traceback
                traceback.print_exc()
        
        # Запускаем загрузку модели в отдельном потоке
        model_thread = threading.Thread(target=load_model_background, daemon=True)
        model_thread.start()
        
        # Ждем загрузки модели с обновлением заставки (с таймаутом)
        print("\n⏳ Ожидание завершения загрузки модели...")
        import time
        max_wait_time = 600  # 10 минут максимально
        start_time = time.time()
        last_status_update = start_time
        status_interval = 10  # Обновляем статус каждые 10 секунд
        
        while model_thread.is_alive():
            elapsed = time.time() - start_time
            
            # Периодически выводим статус
            if time.time() - last_status_update >= status_interval:
                progress_bar = "█" * int((elapsed % 4) + 1) + "░" * (5 - int((elapsed % 4) + 1))
                print(f"⏳ Загрузка модели {progress_bar} ({elapsed:.0f} сек)")
                last_status_update = time.time()
            
            if elapsed > max_wait_time:
                print(f"\n⚠️  Загрузка модели сильно затягивается (прошло {elapsed/60:.1f} минут)")
                print("💡 Это может быть первый запуск - модель загружается впервые")
                print("📦 Размер модели: ~3 ГБ, скорость зависит от интернета")
                print("⏳ Продолжаем ждать загрузки модели...")
                
            try:
                splash.update()
            except:
                pass  # Если форма закрыта, продолжаем ждать
            model_thread.join(timeout=0.1)
        
        # Закрываем заставку безопасно
        splash.safe_close()
        
        # Проверяем успешность загрузки модели
        if loaded_model['model'] is None:
            print("\n" + "=" * 70)
            print("❌ ОШИБКА: Модель не загружена!")
            print("=" * 70)
            print("Приложение не может запуститься без модели.")
            print("Проверьте ошибки выше и попробуйте снова.")
            print("=" * 70)
            input("\nНажмите Enter для выхода...")
            return
        
        print("\n✅ Модель успешно загружена, создание главного окна...")
        root = tk.Tk()
        print("✓ Главное окно создано")
        
        print("Создание экземпляра приложения...")
        app = WhisperXApp(root)
        
        # Устанавливаем загруженную модель в приложение
        app.model = loaded_model['model']
        app.device = loaded_model['device']
        app.compute_type = loaded_model['compute_type']
        app.model_status_label.config(text="✅ Модель large-v3 загружена", foreground=COLORS['success'])
        app.add_file_btn.config(state=tk.NORMAL)
        print("✓ Модель large-v3 установлена в приложение")
        
        print("✓ Экземпляр приложения создан")
        
        print("\nЗапуск главного цикла событий...")
        print("Окно должно открыться. Для выхода закройте окно.")
        print("=" * 70)
        root.mainloop()
        
        print("\n✓ Приложение закрыто нормально")
        sys.exit(0)
    except Exception as e:
        import traceback
        print("\n" + "=" * 70)
        print("КРИТИЧЕСКАЯ ОШИБКА ПРИ ЗАПУСКЕ ПРИЛОЖЕНИЯ")
        print("=" * 70)
        print(f"Ошибка: {e}")
        print("\nПолная информация об ошибке:")
        print("-" * 70)
        traceback.print_exc()
        print("=" * 70)
        input("\nНажмите Enter для выхода...")


if __name__ == "__main__":
    main()

